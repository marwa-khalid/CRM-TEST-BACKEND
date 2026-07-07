import io
import os
import copy
import zipfile
from datetime import date, datetime
from sqlalchemy.orm import Session
from docx import Document
from docx.text.paragraph import Paragraph
import openpyxl

from appflow.logger import logger
from libdata.models.tables import (
    ABIBHRCharges, Claim, ClientDetail, EngineerDetail,
    HireDetail, HireVehicleProvided, InsurerBroker, LocationCondition,
    PlatingAdditionalCharges, Recovery, RouteRepair, Storage, ThirdPartyInsurer, TotalLoss,
)
from libdata.enums import PersonRoleEnum
from appflow.utils import build_case_reference

ASSET_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "assets", "PaymentPackFiles")
)

_TEMPLATES = {
    "front_cover":    "Front Cover Information (1).docx",
    "credit_invoice": "Credit Invoice (1).docx",
    "abi_breakdown":  "Breakdown of ABI Hire Charges  (2).docx",
    "plating":        "Plating Invoice  (1).docx",
    "hire_validation":"Hire Period Validation Form  (1).docx",
    "covering_letter":"Updated Covering Letter.xlsx",
}


# ── helpers ────────────────────────────────────────────────────────────────────

def _f(v) -> float:
    try:
        return float(str(v or 0)) or 0.0
    except (ValueError, TypeError):
        return 0.0


def _money(v: float) -> str:
    return f"£{v:,.2f}"


def _ordinal(d) -> str:
    if d is None:
        return ""
    if isinstance(d, datetime):
        d = d.date()
    day = d.day
    sfx = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return d.strftime(f"%-d{sfx} %B %Y")


def _short(d) -> str:
    if d is None:
        return ""
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d/%m/%Y")


def _short2(d) -> str:
    if d is None:
        return ""
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d/%m/%y")


def _set_para(para, text: str):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _set_cell(cell, text: str):
    if not cell.paragraphs:
        return
    lines = str(text).split("\n")
    first_p = cell.paragraphs[0]
    # Remove any other template paragraphs (avoids blank gap lines)
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    # First line goes into the original (formatted) paragraph
    _set_para(first_p, lines[0])
    # Extra lines clone the first paragraph so every line keeps identical
    # formatting (font/size), one line per vehicle
    prev = first_p._element
    for line in lines[1:]:
        new_el = copy.deepcopy(first_p._element)
        prev.addnext(new_el)
        prev = new_el
        _set_para(Paragraph(new_el, first_p._parent), line)


def _load(key: str) -> Document:
    return Document(os.path.join(ASSET_DIR, _TEMPLATES[key]))


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── document builders ─────────────────────────────────────────────────────────

def _front_cover(insured: str, their_ref: str, policy_no: str, incident_date) -> bytes:
    doc = _load("front_cover")
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Your Insured:"):
            _set_para(p, f"Your Insured: {insured}")
        elif t.startswith("Your Reference:"):
            _set_para(p, f"Your Reference: {their_ref}")
        elif t.startswith("Your Policy Number:"):
            _set_para(p, f"Your Policy Number: {policy_no}")
        elif t.startswith("Date of Incident:"):
            _set_para(p, f"Date of Incident: {_ordinal(incident_date)}")
    return _to_bytes(doc)


def _credit_invoice(
    insurer_company: str, today: date, invoice_no: str,
    hire_start, hire_end, total_days: int,
    client_name: str, vehicle_reg: str, vehicle_desc: str,
    bhr_rate: float, col_del: float, admin_fee: float, cdw: float,
) -> bytes:
    doc = _load("credit_invoice")
    hire_total = bhr_rate * total_days
    cdw_per_day = cdw  # cdw_charges is stored as the per-day rate
    cdw_total = cdw_per_day * total_days  # CDW for BHR = CDW × hire days
    sub = hire_total + col_del + admin_fee + cdw_total
    vat = sub * 0.2
    grand = sub + vat

    paras = doc.paragraphs
    for p in paras:
        t = p.text
        if "Invoice Date:" in t:
            _set_para(p, f"Invoice Date:\t\t{_short2(today)}\t")
        elif "Invoice Number:" in t:
            _set_para(p, f"Invoice Number:  \t{invoice_no}")
        elif "Hire Start:" in t:
            _set_para(p, f"Hire Start:\t\t{_short(hire_start)}")
        elif "Hire End:" in t:
            _set_para(p, f"Hire End:\t\t{_short(hire_end)}")
        elif "Total Hire Days:" in t:
            _set_para(p, f"Total Hire Days: \t{total_days} Days")

    # Bill To section: first non-empty para after "Bill To:"
    found_bill = False
    bill_line = 0
    for p in paras:
        if p.text.strip() == "Bill To:":
            found_bill = True
            bill_line = 0
            continue
        if found_bill and p.text.strip():
            bill_line += 1
            if bill_line == 1:
                _set_para(p, insurer_company)
            elif bill_line == 2:
                _set_para(p, "")
                found_bill = False

    tables = doc.tables
    if tables:
        t0 = tables[0]
        if len(t0.rows) > 1:
            r = t0.rows[1]
            _set_cell(r.cells[0], client_name)
            _set_cell(r.cells[1], vehicle_reg)
            _set_cell(r.cells[2], vehicle_desc)

    if len(tables) > 1:
        t1 = tables[1]
        rows = t1.rows
        # Row 1: hire charge
        if len(rows) > 1:
            _set_cell(rows[1].cells[0], f"{total_days} Days Charged at Basic Hire Rate")
            _set_cell(rows[1].cells[1], str(total_days))
            _set_cell(rows[1].cells[2], _money(bhr_rate))
            _set_cell(rows[1].cells[3], _money(hire_total))
        # Row 2: collection & delivery
        if len(rows) > 2:
            _set_cell(rows[2].cells[3], _money(col_del))
        # Row 3: admin fee
        if len(rows) > 3:
            _set_cell(rows[3].cells[3], _money(admin_fee))
        # Row 4: CDW
        if len(rows) > 4:
            _set_cell(rows[4].cells[1], str(total_days))
            _set_cell(rows[4].cells[2], _money(cdw_per_day))
            _set_cell(rows[4].cells[3], _money(cdw_total))
        # Row 5: Sub Total
        if len(rows) > 5:
            _set_cell(rows[5].cells[-1], _money(sub))
        # Row 6: VAT
        if len(rows) > 6:
            _set_cell(rows[6].cells[-1], _money(vat))
        # Row 7: TOTAL
        if len(rows) > 7:
            _set_cell(rows[7].cells[-1], _money(grand))

    return _to_bytes(doc)


def _abi_breakdown(
    our_ref: str, their_ref: str, hire_start, hire_end,
    total_days: int, vehicle_group: str,
    abi_rate: float, abi_extra: float,
) -> bytes:
    doc = _load("abi_breakdown")
    abi_total_rate = abi_rate + abi_extra
    abi_total = abi_total_rate * total_days

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Our Reference:"):
            _set_para(p, f"Our Reference: {our_ref}")
        elif t.startswith("Your Reference:"):
            _set_para(p, f"Your Reference: {their_ref}")

    if doc.tables:
        rows = doc.tables[0].rows
        for r in rows:
            label = r.cells[0].text.strip() if r.cells else ""
            # Use cells[2] for value (merged cols 2+3)
            val_cell = r.cells[2] if len(r.cells) > 2 else r.cells[-1]
            if "Vehicle Group" in label:
                _set_cell(val_cell, vehicle_group)
            elif "Hire Start Date" in label:
                _set_cell(r.cells[1], _short(hire_start))
                # cells[2] = "Total Days Hired" label, cells[3] = its value
                if len(r.cells) > 3 and "Total Days Hired" in r.cells[2].text:
                    _set_cell(r.cells[3], str(total_days))
            elif "Hire End Date" in label:
                _set_cell(r.cells[1], _short(hire_end))
                if len(r.cells) > 3 and "Total Days Hired" in r.cells[2].text:
                    _set_cell(r.cells[3], str(total_days))
            elif "ABI Hire Rate per day" in label:
                _set_cell(val_cell, _money(abi_rate))
            elif "Total Additional Daily Charges" in label:
                _set_cell(val_cell, _money(abi_extra))
            elif "Total Daily ABI Rate Including Additional" in label:
                _set_cell(val_cell, _money(abi_total_rate))
            elif "Total ABI Costs Including Additional" in label:
                _set_cell(val_cell, _money(abi_total))

    return _to_bytes(doc)


def _plating_invoice(
    insurer_company: str, today: date, invoice_no: str,
    client_name: str, vehicle_reg: str, vehicle_desc: str,
    mot_cost: float, plating_fee: float, plating_total: float,
) -> bytes:
    doc = _load("plating")

    for p in doc.paragraphs:
        t = p.text
        if "Invoice Date:" in t:
            _set_para(p, f"Invoice Date:\t\t{_short2(today)}\t")
        elif "Invoice Number:" in t:
            _set_para(p, f"Invoice Number:  \t{invoice_no}")

    found_bill = False
    bill_line = 0
    for p in doc.paragraphs:
        if p.text.strip() == "Bill To:":
            found_bill = True
            bill_line = 0
            continue
        if found_bill and p.text.strip():
            bill_line += 1
            if bill_line == 1:
                _set_para(p, insurer_company)
            elif bill_line == 2:
                _set_para(p, "")
                found_bill = False

    tables = doc.tables
    if tables:
        t0 = tables[0]
        if len(t0.rows) > 1:
            r = t0.rows[1]
            _set_cell(r.cells[0], client_name)
            _set_cell(r.cells[1], vehicle_reg)
            _set_cell(r.cells[2], vehicle_desc)

    if len(tables) > 1:
        rows = tables[1].rows
        for r in rows:
            label = r.cells[0].text.strip() if r.cells else ""
            val_cell = r.cells[-1]
            if "MOT" in label:
                _set_cell(val_cell, _money(mot_cost))
            elif "Plating Costs" in label:
                _set_cell(val_cell, _money(plating_fee))
            elif label == "Total":
                _set_cell(val_cell, _money(plating_total))

    return _to_bytes(doc)


def _hire_validation(
    notification_date, inspection_date,
    repair_auth, repair_start, repair_completed,
    settlement_offer, offer_accepted, payment_received,
    cil_cheque_received,
) -> bytes:
    doc = _load("hire_validation")
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Date of notification by Insured:"):
            _set_para(p, f"Date of notification by Insured:\t\t\t\t{_short(notification_date)}\t")
        elif t.startswith("Date of inspection:"):
            _set_para(p, f"Date of inspection:\t\t\t\t\t\t{_short(inspection_date)}")
        elif t.startswith("Date Repairs Authorised:"):
            _set_para(p, f"Date Repairs Authorised:\t\t{_short(repair_auth)}")
        elif t.startswith("Date Repairs Started:"):
            _set_para(p, f"Date Repairs Started:\t\t\t{_short(repair_start)}")
        elif t.startswith("Date Satisfaction Note Signed:"):
            _set_para(p, f"Date Satisfaction Note Signed:\t{_short(repair_completed)}")
        elif t.startswith("Date of settlement offer:"):
            _set_para(p, f"Date of settlement offer: \t\t{_short(settlement_offer)}")
        elif t.startswith("Date offer accepted:"):
            _set_para(p, f"Date offer accepted: \t\t{_short(offer_accepted)}")
        elif t.startswith("Date payment received:"):
            _set_para(p, f"Date payment received: \t\t{_short(payment_received)}")
        elif t.startswith("Date CIL cheque received:"):
            _set_para(p, f"Date CIL cheque received:\t\t{_short(cil_cheque_received)}")

    # Blank the "If applicable – explanation for delays" header and its sample
    # lines (keep the empty paragraphs so the downloader has space to write)
    clear_mode = False
    for p in doc.paragraphs:
        if p.text.strip().startswith("If applicable"):
            for run in p.runs:
                run.text = ""
            clear_mode = True
            continue
        if clear_mode and p.text.strip():
            for run in p.runs:
                run.text = ""
    return _to_bytes(doc)


def _covering_letter(
    insurer_company: str, today: date, insured_name: str,
    their_ref: str, client_name: str, our_ref: str,
    incident_date, notification_date,
    bhr_hire: float, abi_hire_30: float, abi_hire_31_60: float, abi_hire_61plus: float,
    bhr_admin: float, abi_admin: float, abi_admin_31_60: float, abi_admin_61plus: float,
    repair: float, storage: float, recovery: float, plating: float,
    engineer_fee: float, cdw: float, col_del: float,
    sign_off_name: str = "",
) -> bytes:
    wb = openpyxl.load_workbook(os.path.join(ASSET_DIR, _TEMPLATES["covering_letter"]))
    ws = wb.active

    # Header fields
    ws["B7"] = insurer_company
    ws["B15"] = _ordinal(today)
    ws["C17"] = insured_name
    ws["C19"] = client_name
    ws["C20"] = our_ref
    ws["C21"] = incident_date if incident_date is None else (
        incident_date.date() if isinstance(incident_date, datetime) else incident_date
    )
    ws["B26"] = notification_date if notification_date is None else (
        notification_date.date() if isinstance(notification_date, datetime) else notification_date
    )

    # Financial data rows  (C=BHR, D=ABI30, E=ABI31-60, F=ABI61+)
    ws["C33"] = round(bhr_hire, 2)
    ws["D33"] = round(abi_hire_30, 2)
    ws["E33"] = round(abi_hire_31_60, 2)
    ws["F33"] = round(abi_hire_61plus, 2)

    ws["C34"] = round(bhr_admin, 2)
    ws["D34"] = round(abi_admin, 2)
    ws["E34"] = round(abi_admin_31_60, 2)
    ws["F34"] = round(abi_admin_61plus, 2)

    # Repair, storage, recovery, plating, engineer are NOT topped up by the ABI
    # band — same value in every column (only hire & admin get the 10/20 uplift)
    for row, base in (("35", repair), ("36", storage), ("37", recovery),
                      ("38", plating), ("39", engineer_fee)):
        val = round(base, 2)
        ws[f"C{row}"] = val
        ws[f"D{row}"] = val
        ws[f"E{row}"] = val
        ws[f"F{row}"] = val

    ws["C40"] = round(cdw, 2)
    ws["D40"] = 0
    ws["E40"] = 0
    ws["F40"] = 0

    ws["C41"] = round(col_del, 2)
    ws["D41"] = 0
    ws["E41"] = 0
    ws["F41"] = 0

    # Sign-off name (replaces the template's hardcoded "Alex Berwick")
    if sign_off_name:
        ws["B58"] = sign_off_name

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ── main entry point ──────────────────────────────────────────────────────────

def generate_payment_pack(claim_id: int, tenant_id: int, db: Session, sign_off_name: str = "") -> tuple:
    today = date.today()

    claim = db.query(Claim).filter(Claim.id == claim_id, Claim.tenant_id == tenant_id).first()
    if not claim:
        raise ValueError(f"Claim {claim_id} not found")

    insurer = db.query(InsurerBroker).filter(InsurerBroker.claim_id == claim_id).first()
    # Third-party people are also ClientDetail rows on this claim — filter to the
    # actual client (driver), not a third party / handling agent.
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT,
    ).first()
    hire_vps = (
        db.query(HireVehicleProvided)
        .filter(
            HireVehicleProvided.claim_id == claim_id,
            HireVehicleProvided.is_active == True,
            HireVehicleProvided.is_deleted == False,
        )
        .order_by(HireVehicleProvided.id.asc())
        .all()
    )
    hire_vp = hire_vps[0] if hire_vps else None
    hire_d = (
        db.query(HireDetail)
        .filter(HireDetail.hire_vehicle_provided_id == hire_vp.id)
        .first()
        if hire_vp else None
    )
    accident = db.query(LocationCondition).filter(LocationCondition.claim_id == claim_id).first()
    engineer = db.query(EngineerDetail).filter(EngineerDetail.claim_id == claim_id).first()
    plating = db.query(PlatingAdditionalCharges).filter(PlatingAdditionalCharges.claim_id == claim_id).first()
    abi_rec = db.query(ABIBHRCharges).filter(ABIBHRCharges.claim_id == claim_id).first()
    repair = db.query(RouteRepair).filter(RouteRepair.claim_id == claim_id).first()
    total_loss = db.query(TotalLoss).filter(TotalLoss.claim_id == claim_id).first()
    storages = db.query(Storage).filter(Storage.claim_id == claim_id).all()
    recoveries = db.query(Recovery).filter(Recovery.claim_id == claim_id).all()

    # ── derive values ──────────────────────────────────────────────────────────
    invoice_no = (abi_rec.invoice_number if abi_rec else None) or ""
    insurer_company = (insurer.company_name if insurer else None) or ""
    tpi = db.query(ThirdPartyInsurer).filter(ThirdPartyInsurer.claim_id == claim_id).first()
    their_ref = (tpi.insurer_reference if tpi else None) or ""  # "Your Reference" = TPI insurer reference
    policy_no = (insurer.policy_number if insurer else None) or ""
    insured_name = (insurer.policy_holder if insurer else None) or ""

    client_name = " ".join(filter(None, [
        getattr(client, "first_name", None),
        getattr(client, "surname", None),
    ])) if client else ""

    # Hire period spans the whole switch sequence: start = first vehicle's start,
    # end = last vehicle's end. This keeps the date range consistent with
    # total_days (summed across all vehicles below).
    hire_start = getattr(hire_vp, "hire_start_date", None) if hire_vp else None
    hire_end = getattr(hire_vps[-1], "hire_end_date", None) if hire_vps else None
    # List every active hire vehicle (one line per car)
    vehicle_reg = "\n".join((v.hire_vehicle_registration or "") for v in hire_vps) if hire_vps else ""
    vehicle_desc = "\n".join(
        f"{v.make or ''} / {v.model or ''}".strip(" /") for v in hire_vps
    ) if hire_vps else ""

    vehicle_group = ""
    if hire_vp and hire_vp.actual_vehicle_category_id:
        from libdata.models.tables import ActualVehicleCategory
        cat = db.query(ActualVehicleCategory).filter(
            ActualVehicleCategory.id == hire_vp.actual_vehicle_category_id
        ).first()
        vehicle_group = cat.label if cat else ""

    abi_rate = _f(getattr(hire_d, "abi_hire_charge_per_day", 0) if hire_d else 0)
    abi_extra = _f(getattr(hire_d, "abi_extra_charges_per_day", 0) if hire_d else 0)
    abi_admin = _f(getattr(hire_d, "abi_administration_fee", 0) if hire_d else 0)
    bhr_admin = _f(getattr(hire_d, "bhr_administration_fee", 0) if hire_d else 0)
    bhr_rate = _f(getattr(hire_d, "bhr_hire_charge_per_day", 0) if hire_d else 0)
    if not bhr_rate:
        bhr_rate = _f(getattr(hire_vp, "rate", 0) if hire_vp else 0)
    # Total days = sum of every active vehicle's total hire days.
    # Total hire = sum of each vehicle's (its own days × its own rate), since
    # different vehicle categories have different daily rates.
    total_days_f = 0.0
    abi_hire_sum = 0.0
    bhr_hire_sum = 0.0
    for vp in hire_vps:
        det = (
            db.query(HireDetail)
            .filter(HireDetail.hire_vehicle_provided_id == vp.id)
            .first()
        )
        if det:
            d = _f(getattr(det, "final_total_no_of_hire_days", 0))
            if not d:
                d = _f(getattr(det, "no_of_days_hire_so_far", 0))
            total_days_f += d
            abi_r = _f(getattr(det, "abi_hire_charge_per_day", 0))
            bhr_r = _f(getattr(det, "bhr_hire_charge_per_day", 0)) or _f(getattr(vp, "rate", 0))
            abi_hire_sum += d * abi_r
            bhr_hire_sum += d * bhr_r
    total_days = int(total_days_f)
    # Blend the per-day rates so the existing `rate × total_days` maths yields the
    # summed hire amount across all vehicles (admin/other charges stay added once).
    if total_days_f > 0:
        abi_rate = abi_hire_sum / total_days_f
        bhr_rate = bhr_hire_sum / total_days_f
    cdw = _f(getattr(hire_d, "cdw_charges", 0) if hire_d else 0)
    col_del = _f(getattr(hire_d, "collection_delivery_fee", 0) if hire_d else 0)

    incident_date = accident.date_time if accident else None
    inspection_date = engineer.inspection_date if engineer else None
    notification_date = getattr(hire_vp, "inst_fleet_on_hire", None) if hire_vp else None

    storage_total = sum(_f(s.total_storage_charges) for s in storages)
    recovery_total = sum(_f(r.recovery_charges) for r in recoveries)
    plating_total = _f(plating.total_plating_cost if plating else 0)
    plating_mot = _f(plating.private_hire_mot_cost if plating else 0)
    plating_fee_val = _f(plating.private_hire_plating_fee if plating else 0)
    engineer_fee = _f(engineer.engineer_fee if engineer else 0)
    repair_cost = _f(repair.sub_total if repair else 0)  # exclusive of VAT

    our_ref = build_case_reference(claim_id, db)  # "Our Reference" — the claim reference

    abi_hire_30 = abi_rate * total_days          # basic ABI rate × days
    abi_hire_31_60 = abi_hire_30 * 1.1           # +10%
    abi_hire_61plus = abi_hire_30 * 1.2          # +20%
    abi_admin_31_60 = abi_admin * 1.1
    abi_admin_61plus = abi_admin * 1.2

    # ── save payment_pack_raised_date + auto invoice reference ─────────────────
    from appflow.utils import build_invoice_reference
    if abi_rec:
        if not abi_rec.payment_pack_raised_date:
            abi_rec.payment_pack_raised_date = today
        if not abi_rec.invoice_number:
            abi_rec.invoice_number = build_invoice_reference(claim_id)
        db.commit()
    else:
        abi_rec = ABIBHRCharges(
            claim_id=claim_id,
            payment_pack_raised_date=today,
            invoice_number=build_invoice_reference(claim_id),
        )
        db.add(abi_rec)
        db.commit()

    # ── build ZIP ──────────────────────────────────────────────────────────────
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("1_Front_Cover_Information.docx",
                    _front_cover(insured_name, their_ref, policy_no, incident_date))
        zf.writestr("2_Credit_Invoice.docx",
                    _credit_invoice(insurer_company, today, invoice_no,
                                    hire_start, hire_end, total_days,
                                    client_name, vehicle_reg, vehicle_desc,
                                    bhr_rate, col_del, bhr_admin, cdw))
        zf.writestr("3_ABI_Hire_Breakdown.docx",
                    _abi_breakdown(our_ref, their_ref, hire_start, hire_end,
                                   total_days, vehicle_group, abi_rate, 0))  # additional charges = 0 for now
        zf.writestr("4_Plating_Invoice.docx",
                    _plating_invoice(insurer_company, today, invoice_no,
                                     client_name, vehicle_reg, vehicle_desc,
                                     plating_mot, plating_fee_val, plating_total))
        zf.writestr("5_Hire_Period_Validation_Form.docx",
                    _hire_validation(
                        notification_date, inspection_date,
                        repair_auth=getattr(repair, "repair_auth", None) if repair else None,
                        repair_start=getattr(repair, "repair_start", None) if repair else None,
                        repair_completed=getattr(repair, "repair_completed", None) if repair else None,
                        settlement_offer=getattr(total_loss, "pav_offer_made_client", None) if total_loss else None,
                        offer_accepted=getattr(total_loss, "pav_offer_accepted", None) if total_loss else None,
                        payment_received=getattr(total_loss, "pav_cheque_received", None) if total_loss else None,
                        cil_cheque_received=getattr(repair, "cil_cheque_request", None) if repair else None,
                    ))
        zf.writestr("6_Covering_Letter.xlsx",
                    _covering_letter(insurer_company, today, insured_name,
                                     their_ref, client_name, our_ref,
                                     incident_date, notification_date,
                                     bhr_rate * total_days, abi_hire_30,
                                     abi_hire_31_60, abi_hire_61plus,
                                     bhr_admin, abi_admin, abi_admin_31_60, abi_admin_61plus,
                                     repair_cost, storage_total, recovery_total,
                                     plating_total, engineer_fee, cdw * total_days, col_del,
                                     sign_off_name))

    zip_buf.seek(0)
    filename = f"PaymentPack_Claim{claim_id}_{today.strftime('%Y%m%d')}.zip"
    return zip_buf.read(), filename
