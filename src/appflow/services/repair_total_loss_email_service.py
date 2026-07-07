import os
import base64
from io import BytesIO
from datetime import datetime
from fastapi import HTTPException
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, ReplyTo, Attachment, FileContent, FileName, FileType, Disposition, ContentId
)
from docx import Document
from sqlalchemy.orm import Session, joinedload

from libdata.models.tables import (
    ClientDetail,
    Claim,
    Address,
    VehicleDetail,
    LocationCondition,
    Handler,
    Referrer,
    HireVehicleProvided,
    EngineerDetail,
    ThirdPartyInsurer,
    RouteRepair,
    CaseDocument,
)
from appflow.services.s3_service import S3Service
from libdata.enums import HistoryLogType, PersonRoleEnum
from appflow.services.history_activity_service import HistoryActivityService
from appflow.utils import actor_id, get_tenant_id, build_case_reference, handler_name_for_claim, handler_name_for_user


BASE_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")

# Logo embedded as CID inline attachment — data URIs are blocked by Gmail/Outlook web.
_PNG_LOGO_PATH = os.path.join(BASE_TEMPLATE_DIR, "logo.png")
LOGO_CID = "na_logo"
LOGO_PNG_B64 = ""
if os.path.exists(_PNG_LOGO_PATH):
    with open(_PNG_LOGO_PATH, "rb") as _f:
        LOGO_PNG_B64 = base64.b64encode(_f.read()).decode()


FROM_EMAIL = os.getenv("SENDGRID_SENDER", "noreplynationwideassist@yopmail.com")
REPLY_TO_EMAIL = os.getenv("SENDGRID_REPLY_TO", FROM_EMAIL)


def _format_date(value, fmt="%d/%m/%Y"):
    if not value:
        return ""
    try:
        return value.strftime(fmt)
    except Exception:
        return str(value)


def _safe(value, fallback="N/A"):
    if value is None or value == "":
        return fallback
    return str(value)


def _get_claim_context(db: Session, claim_id: int):
    claim = (
        db.query(Claim)
        .options(joinedload(Claim.claim_type))
        .filter(Claim.id == claim_id)
        .first()
    )
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    client = (
        db.query(ClientDetail)
        .filter(
            ClientDetail.claim_id == claim_id,
            ClientDetail.role == PersonRoleEnum.CLIENT,
        )
        .first()
    )
    if not client:
        client = (
            db.query(ClientDetail)
            .filter(ClientDetail.claim_id == claim_id, ClientDetail.role == "CLIENT")
            .first()
        )
    if not client:
        client = (
            db.query(ClientDetail)
            .filter(ClientDetail.claim_id == claim_id)
            .first()
        )

    client_address = (
        db.query(Address)
        .filter(Address.id == client.address_id)
        .first()
        if client and client.address_id
        else None
    )

    vehicle = (
        db.query(VehicleDetail)
        .filter(VehicleDetail.claim_id == claim_id)
        .first()
    )

    incident = (
        db.query(LocationCondition)
        .filter(LocationCondition.claim_id == claim_id)
        .first()
    )

    referrer = (
        db.query(Referrer)
        .filter(Referrer.claim_id == claim_id)
        .first()
    )

    handler = (
        db.query(Handler)
        .filter(Handler.id == claim.handler_id)
        .first()
        if claim and claim.handler_id
        else None
    )

    latest_hire = (
        db.query(HireVehicleProvided)
        .filter(
            HireVehicleProvided.claim_id == claim_id,
            HireVehicleProvided.is_active == True,
            HireVehicleProvided.is_deleted == False,
        )
        .order_by(HireVehicleProvided.id.desc())
        .first()
    )

    tpi = (
        db.query(ThirdPartyInsurer)
        .filter(ThirdPartyInsurer.claim_id == claim_id)
        .first()
    )

    route_repair = (
        db.query(RouteRepair)
        .filter(RouteRepair.claim_id == claim_id)
        .first()
    )

    reference = build_case_reference(claim_id, db)

    client_name = ""
    salutation = ""
    if client:
        client_name = f"{client.first_name or ''} {client.surname or ''}".strip()
        gender = (client.gender or "").lower()
        title = "Mr" if gender == "male" else "Mrs" if gender == "female" else "Mr/Mrs"
        salutation = f"{title} {client.surname or ''}".strip()

    return {
        "claim": claim,
        "client": client,
        "client_address": client_address,
        "vehicle": vehicle,
        "incident": incident,
        "referrer": referrer,
        "handler": handler,
        "latest_hire": latest_hire,
        "tpi": tpi,
        "route_repair": route_repair,
        "reference": reference,
        "client_name": client_name or "Client",
        "salutation": salutation or "Client",
    }


def _resolve_recipient(data_to_email, fallback_email):
    recipient = data_to_email or fallback_email
    if not recipient:
        raise HTTPException(status_code=400, detail="Recipient email is missing")
    return recipient


def _split_recipients(to_email: str):
    recipients = []
    seen = set()
    for email in str(to_email or "").replace(",", ";").split(";"):
        clean_email = email.strip()
        key = clean_email.lower()
        if clean_email and key not in seen:
            recipients.append(clean_email)
            seen.add(key)
    return recipients


def _configured_copy_recipients():
    return ";".join(
        [
            os.getenv("HIRE_INSTRUCTION_COPY_EMAIL", ""),
            os.getenv("MS_GRAPH_MAILBOX", ""),
            os.getenv("OUTLOOK_MAILBOX", ""),
        ]
    )


def _send_email(to_email: str, subject: str, html_content: str, copy_to: str = "", attachments=None):
    recipients = _split_recipients(";".join([str(to_email or ""), str(copy_to or "")]))
    if not recipients:
        raise HTTPException(status_code=400, detail="Recipient email is missing")

    inline_logo = (
        [{"cid": LOGO_CID, "content_bytes": LOGO_PNG_B64, "content_type": "image/png"}]
        if LOGO_PNG_B64 else None
    )

    # Prefer Microsoft Graph (delivered from a real Outlook mailbox); SendGrid
    # from the unverified yopmail.com sender is silently dropped by Outlook.
    from appflow.services.graph_email_service import GraphEmailService
    if GraphEmailService.is_configured():
        result = GraphEmailService.send_mail(
            recipients, subject, html_content,
            reply_to=REPLY_TO_EMAIL,
            inline_images=inline_logo,
            attachments=attachments,
        )
        if result is not None:
            return result.status_code

    api_key = os.getenv("SENDGRID_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="SendGrid API key not found")

    message = Mail(
        from_email=FROM_EMAIL,
        to_emails=recipients,
        subject=subject,
        html_content=html_content,
    )
    message.reply_to = ReplyTo(REPLY_TO_EMAIL, "No-Reply")

    if LOGO_PNG_B64:
        attachment = Attachment(
            FileContent(LOGO_PNG_B64),
            FileName("logo.png"),
            FileType("image/png"),
            Disposition("inline"),
            ContentId(LOGO_CID),
        )
        message.add_attachment(attachment)

    for att in attachments or []:
        message.add_attachment(
            Attachment(
                FileContent(att["content_bytes"]),
                FileName(att.get("name") or "attachment"),
                FileType(att.get("content_type") or "application/octet-stream"),
                Disposition("attachment"),
            )
        )

    try:
        sg = SendGridAPIClient(api_key)
        response = sg.send(message)
        return response.status_code
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"SendGrid send error: {str(e)}")


def _log_email_activity(db, request, claim_id: int, file_name: str, file_type: HistoryLogType):
    HistoryActivityService.create_activity(
        db=db,
        claim_id=claim_id,
        file_name=file_name,
        file_path="",
        file_type=file_type,
        user_id=actor_id(request),
        tenant_id=get_tenant_id(request),
    )


def _row(label, value):
    return f"""
    <div style="display:flex; align-items:center; padding:8px 0;">
        <div style="width:130px; color:#334155; font-size:12px; font-weight:400; font-family:Arial, sans-serif;">{label}</div>
        <div style="color:#334155; font-size:12px; font-weight:600; font-family:Arial, sans-serif;">{_safe(value)}</div>
    </div>
    """


def _divider():
    return '<div style="height:1px; background-color:#e2e8f0; width:100%;"></div>'


def instruct_fleet_off_hire_email(db: Session, claim_id: int, request, to_email: str = None):
    ctx = _get_claim_context(db, claim_id)

    latest_hire = ctx["latest_hire"]
    client_address = ctx["client_address"]

    fleet_email = to_email or "ayesha.rana@nationwideassist.co.uk;ayesha.rana@nationwideassist.co.uk"
    reference = ctx["reference"]

    hire_vehicle = _safe(latest_hire.hire_vehicle_registration if latest_hire else None)
    cl_mobile = _safe(client_address.mobile_tel if client_address else None)
    referrer = ctx["referrer"]
    referrer_name = _safe(
        (referrer.company_name or referrer.contact_name) if referrer else None
    )

    # Use hire end date if available, otherwise today
    if latest_hire and latest_hire.hire_end_date:
        off_hire_date = _format_date(latest_hire.hire_end_date)
    else:
        off_hire_date = datetime.now().strftime("%d/%m/%Y")

    subject = f"New Instruction to Fleet to Off Hire Vehicle (CIL) - {reference}"

    def _info_row(label, value):
        return f"""
        <tr>
          <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400;
                     color:#444444; font-family:Arial,sans-serif;">{label}</td>
          <td style="padding:8px 0; font-size:12px; font-weight:600;
                     color:#444444; font-family:Arial,sans-serif;">{value}</td>
        </tr>
        <tr><td colspan="2" style="padding:0; height:1px; background:#CCCCCC;"></td></tr>
        """

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <body style="margin:0; padding:0; background:#ffffff;">
    <table width="640" cellpadding="0" cellspacing="0" border="0"
           style="margin:0 auto; background:#ffffff; font-family:Arial,sans-serif;">
      <tr>
        <td style="padding:40px 30px 24px 30px; text-align:center;">
          <img src="cid:na_logo" width="33" height="31" alt="Logo"
               style="display:inline-block;">
        </td>
      </tr>

      <!-- Info card -->
      <tr>
        <td style="padding:0 128px 24px 128px;">
          <table width="384" cellpadding="0" cellspacing="0" border="0"
                 style="border:1px solid #CCCCCC; border-radius:8px;
                        padding:0 16px; background:#ffffff;">
            <tr><td style="padding:8px 0 0 0;"></td></tr>
            {_info_row("Reference", _safe(reference))}
            {_info_row("Referrer", referrer_name)}
            {_info_row("Client", _safe(ctx["client_name"]))}
            {_info_row("Hire Vehicle", hire_vehicle)}
            <tr>
              <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400;
                         color:#444444; font-family:Arial,sans-serif;">Cl Mobile No</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600;
                         color:#444444; font-family:Arial,sans-serif;">{cl_mobile}</td>
            </tr>
            <tr><td style="padding:0 0 8px 0;"></td></tr>
          </table>
        </td>
      </tr>

      <!-- Message -->
      <tr>
        <td style="padding:0 160px 40px 160px; text-align:center;">
          <p style="margin:0; font-size:14px; font-weight:400; color:#444444;
                    line-height:1.6; font-family:Arial,sans-serif;">
            Hi,<br><br>
            Please contact the Client to arrange the off hire of this vehicle for
            <strong style="font-weight:600;">{off_hire_date}</strong>.
          </p>
        </td>
      </tr>

      <!-- Divider -->
      <tr>
        <td style="padding:0 30px;">
          <div style="height:1px; background:#CCCCCC; width:580px;"></div>
        </td>
      </tr>

      <!-- Footer -->
      <tr>
        <td style="padding:24px 30px 40px 30px; text-align:center;">
          <p style="margin:0 0 4px 0; font-size:12px; font-weight:600;
                    color:#444444; font-family:Arial,sans-serif;">Kind regards,</p>
          <p style="margin:0; font-size:14px; font-weight:600;
                    color:#444444; font-family:Arial,sans-serif;">
            Nationwide Assist IT / Systems Team
          </p>
        </td>
      </tr>
    </table>
    </body>
    </html>
    """

    status_code = _send_email(
        fleet_email,
        subject,
        html_content,
        copy_to=_configured_copy_recipients(),
    )

    _log_email_activity(
        db=db,
        request=request,
        claim_id=claim_id,
        file_name=f"{subject} for claim {reference}",
        file_type=HistoryLogType.FLEET_OFF_HIRE_INSTRUCTED,
    )

    return {"status": "success", "sendgrid_status": status_code}


def send_cil_to_client_email(db: Session, claim_id: int, request, to_email: str = None):
    ctx = _get_claim_context(db, claim_id)
    client_address = ctx["client_address"]
    vehicle = ctx["vehicle"]
    incident = ctx["incident"]

    receiver_email = _resolve_recipient(
        to_email,
        client_address.email if client_address else None,
    )

    reference = ctx["reference"]
    today = datetime.now().strftime("%-d %B %Y")
    handler_name = handler_name_for_user(db, actor_id(request)) or handler_name_for_claim(ctx["claim"], db) or (ctx["handler"].label if ctx["handler"] else "Claims Handler")
    vehicle_reg = _safe(vehicle.registration if vehicle else None)
    incident_date = _format_date(incident.date_time) if incident else "N/A"

    addr_lines = []
    if client_address:
        if client_address.address:
            addr_lines.append(client_address.address)
        if client_address.postcode:
            addr_lines.append(client_address.postcode)
    addr_block = "<br>".join(addr_lines)

    subject = f"Cash in Lieu Payment - {reference}"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <body style="margin:0; padding:0; background:#ffffff;">
    <table width="640" cellpadding="0" cellspacing="0" border="0"
           style="margin:0 auto; background:#ffffff; font-family:Arial,sans-serif; color:#444444;">

      <!-- Logo -->
      <tr>
        <td style="padding:40px 30px 30px 30px; text-align:center;">
          <img src="cid:na_logo" width="48" alt="Logo" style="display:inline-block;height:auto;">
        </td>
      </tr>

      <!-- Client address block -->
      <tr>
        <td style="padding:0 30px 20px 30px; font-size:14px; line-height:1.6;">
          <strong>{ctx["salutation"]}</strong><br>
          {addr_block}
        </td>
      </tr>

      <!-- Our ref + date -->
      <tr>
        <td style="padding:0 30px 20px 30px; font-size:14px; line-height:1.6;">
          Our Ref: {_safe(reference)}<br>
          {today}
        </td>
      </tr>

      <!-- Info card -->
      <tr>
        <td style="padding:0 30px 24px 30px;">
          <table width="384" cellpadding="0" cellspacing="0" border="0"
                 style="border:1px solid #CCCCCC; border-radius:8px; padding:0 16px; background:#ffffff;">
            <tr><td style="padding:8px 0 0 0;"></td></tr>
            <tr>
              <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Incident Date</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{incident_date}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Vehicle</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{vehicle_reg}</td>
            </tr>
            <tr><td style="padding:0 0 8px 0;"></td></tr>
          </table>
        </td>
      </tr>

      <!-- Body -->
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:14px; line-height:1.6;">
          <strong>Dear {ctx["salutation"]},</strong>
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:14px; line-height:1.6;">
          We are pleased to enclose your cheque in respect of the Cash in Lieu
          for the repairs to the above vehicle.
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 40px 30px; font-size:14px; line-height:1.6;">
          If you have any questions or queries, please do not hesitate to contact
          us on the number provided.
        </td>
      </tr>

      <!-- Sign-off -->
      <tr>
        <td style="padding:0 30px 40px 30px; font-size:14px; line-height:1.8;">
          <strong>{handler_name}</strong><br>
          Claims Handler<br>
          Nationwide Assist Ltd
        </td>
      </tr>

      <!-- Divider -->
      <tr>
        <td style="padding:0 30px;">
          <div style="height:1px; background:#CCCCCC; width:580px;"></div>
        </td>
      </tr>

      <!-- Footer -->
      <tr>
        <td style="padding:24px 30px 40px 30px; text-align:center;">
          <p style="margin:0 0 4px 0; font-size:12px; font-weight:600; color:#444444;">Kind regards,</p>
          <p style="margin:0; font-size:14px; font-weight:600; color:#444444;">
            Nationwide Assist IT / Systems Team
          </p>
        </td>
      </tr>
    </table>
    </body>
    </html>
    """

    status_code = _send_email(receiver_email, subject, html_content)

    _log_email_activity(
        db=db,
        request=request,
        claim_id=claim_id,
        file_name=f"CIL sent to client for claim {reference}",
        file_type=HistoryLogType.CIL_SENT_TO_CLIENT,
    )

    return {"status": "success", "sendgrid_status": status_code}


def send_cil_agreement_email(db: Session, claim_id: int, request, to_email: str = None):
    ctx = _get_claim_context(db, claim_id)
    client_address = ctx["client_address"]
    route_repair = ctx["route_repair"]

    receiver_email = _resolve_recipient(
        to_email,
        client_address.email if client_address else None,
    )

    reference = ctx["reference"]
    today = datetime.now().strftime("%-d %B %Y")

    # Fee amounts from route repair record
    fee_excl_vat = f"£{route_repair.net_cil_amount or '0.00'}" if route_repair else "£0.00"
    fee_charged = f"£{route_repair.total_inc_vat or '0.00'}" if route_repair else "£0.00"

    addr_lines = [ctx["salutation"]]
    if client_address:
        if client_address.address:
            addr_lines.append(client_address.address)
        if client_address.postcode:
            addr_lines.append(client_address.postcode)
    addr_block = "<br>".join(addr_lines)

    subject = f"CIL Agreement - {reference}"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <body style="margin:0; padding:0; background:#ffffff;">
    <table width="640" cellpadding="0" cellspacing="0" border="0"
           style="margin:0 auto; background:#ffffff; font-family:Arial,sans-serif; color:#444444;">

      <!-- Logo -->
      <tr>
        <td style="padding:40px 30px 30px 30px; text-align:center;">
          <img src="cid:na_logo" width="48" alt="Logo" style="display:inline-block;height:auto;">
        </td>
      </tr>

      <!-- Header -->
      <tr>
        <td style="padding:0 30px 24px 30px; text-align:center; font-size:18px;
                   font-weight:700; color:#000000; letter-spacing:0.5px;">
          CASH IN LIEU AGREEMENT
        </td>
      </tr>

      <!-- Agreement details table -->
      <tr>
        <td style="padding:0 30px 24px 30px;">
          <table width="580" cellpadding="0" cellspacing="0" border="0"
                 style="border:1px solid #CCCCCC; border-radius:8px; padding:0 16px; background:#ffffff;">
            <tr><td style="padding:8px 0 0 0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Client Name</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{ctx["salutation"]}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Client Address</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444; line-height:1.5;">{addr_block}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Claim Number</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{_safe(reference)}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Date</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{today}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Service</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">Recovery of Cash In Lieu</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:160px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Fee Charged</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{fee_charged} + VAT</td>
            </tr>
            <tr><td style="padding:0 0 8px 0;"></td></tr>
          </table>
        </td>
      </tr>

      <!-- Agreement body -->
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:13px; line-height:1.7; color:#444444;">
          I confirm that I <strong>{ctx["client_name"]}</strong> have requested CAMS Ltd/Nationwide Assist
          to recover a Cash in Lieu payment in respect of my vehicle damages in the amount of
          <strong>{fee_excl_vat}</strong> exclusive of VAT. For this service I agree to pay a fixed fee
          of <strong>{fee_charged} + VAT</strong>.
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:13px; line-height:1.7; color:#444444;">
          Where this agreement is made during a visit to your or another person's home, or to your place
          of work, or during an excursion arranged by us, or after an offer made by you during such a
          visit or excursion, and where in making this contract you are acting for purposes which are
          outside your trade or profession, you have the right to cancel the contract within a period of
          7 days starting with the date on which you receive this notice by delivering or sending
          (including by electronic mail) a cancellation notice to us to the following address:
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:13px; line-height:1.7; color:#444444;">
          <strong>By post to:</strong><br>
          Cancellations Team, CAMS LTD / Nationwide Assist,<br>
          25-39 Small Heath Highway, Birmingham B10 0EU<br><br>
          <strong>Or by Email to:</strong> claims@camsna.co.uk
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:13px; line-height:1.7; color:#444444;">
          Notice of cancellation will be deemed to be served on the day it is posted or sent to us,
          or if it is sent to us by electronic mail, the day it is sent.
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 32px 30px; font-size:13px; line-height:1.7; color:#444444;">
          If you wish to cancel this service after the 7 day period, you will still be liable for the
          <strong>{fee_charged} + VAT</strong> administration charge.
        </td>
      </tr>

      <!-- Signature block -->
      <tr>
        <td style="padding:0 30px 40px 30px;">
          <table width="580" cellpadding="0" cellspacing="0" border="0">
            <tr>
              <td width="185" style="font-size:13px; color:#444444; padding-right:20px;">
                Signed..........................................................
              </td>
              <td width="185" style="font-size:13px; color:#444444; padding-right:20px;">
                Print name......................................................
              </td>
              <td width="185" style="font-size:13px; color:#444444;">
                Date................................................................
              </td>
            </tr>
          </table>
        </td>
      </tr>

      <!-- Divider -->
      <tr>
        <td style="padding:0 30px;">
          <div style="height:1px; background:#CCCCCC; width:580px;"></div>
        </td>
      </tr>

      <!-- Company footer -->
      <tr>
        <td style="padding:20px 30px 40px 30px; text-align:center;
                   font-size:12px; color:#444444; line-height:1.6;">
          CAMS LTD / Nationwide Assist<br>
          25-39 Small Heath Highway, Birmingham, B10 0EU<br>
          T: 0121 7667515 &nbsp;|&nbsp; F: 0121 7660489<br>
          E: claims@camsna.co.uk &nbsp;|&nbsp; W: www.nationwideassist.co.uk
        </td>
      </tr>
    </table>
    </body>
    </html>
    """

    status_code = _send_email(receiver_email, subject, html_content)

    _log_email_activity(
        db=db,
        request=request,
        claim_id=claim_id,
        file_name=f"CIL agreement sent for claim {reference}",
        file_type=HistoryLogType.CIL_AGREEMENT_SENT,
    )

    return {"status": "success", "sendgrid_status": status_code}


def _build_tpi_cover_letter_docx(ctx) -> bytes:
    """Covering letter that goes with the engineer's report to the TPI —
    'Letter Enclosing Copy of Engineers Report for Repair Authority', populated
    from the claim (the legacy .DOC template is unreadable, so it's built here)."""
    claim = ctx["claim"]
    vehicle = ctx["vehicle"]
    incident = ctx["incident"]
    tpi = ctx["tpi"]
    reference = ctx["reference"]
    handler = ctx.get("resolved_handler") or handler_name_for_claim(claim) or (ctx["handler"].label if ctx["handler"] else "Claims Handler")

    veh = " / ".join(filter(None, [
        (vehicle.registration or "") if vehicle else "",
        " ".join(filter(None, [
            (vehicle.make or "") if vehicle else "",
            (vehicle.model or "") if vehicle else "",
        ])).strip(),
    ]))
    incident_date = incident.date_time.strftime("%d/%m/%Y") if incident and incident.date_time else ""

    doc = Document()

    def line(text="", bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        return para

    # Recipient (TPI)
    if tpi:
        for part in [
            getattr(tpi, "company_name", "") or getattr(tpi, "name", "") or "",
            getattr(tpi, "address", "") or "",
            getattr(tpi, "postcode", "") or "",
        ]:
            if part:
                line(str(part))

    line(datetime.now().strftime("%d/%m/%Y"))
    line(f"Our Ref: {reference}")
    line()
    line(f"Client: {ctx['client_name']}", bold=True)
    if veh:
        line(f"Vehicle: {veh}", bold=True)
    if incident_date:
        line(f"Incident Date: {incident_date}", bold=True)
    line()
    line("Dear Sirs")
    line()
    line("Please find enclosed a copy of the Independent Engineer's Report in respect of the above claim.")
    line()
    line("We are seeking your authority to proceed with the repairs as detailed in the report. "
         "Please confirm your repair authority at your earliest convenience.")
    line()
    line("Should you have any questions or queries please do not hesitate to contact us on the number provided.")
    line()
    line("Yours faithfully")
    line()
    line(handler, bold=True)
    line("Claims Handler")
    line("Nationwide Assist Ltd")
    line("T: 0121 766 7515")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _collect_tpi_report_attachments(db: Session, claim_id: int, ctx) -> list:
    """The two attachments the TPI email promises: the uploaded Independent
    Engineer's Report + our covering letter."""
    attachments: list = []
    s3 = S3Service()

    # 1) The engineer's report the handler uploaded on the Engineer Details screen.
    report_doc = (
        db.query(CaseDocument)
        .filter(
            CaseDocument.claim_id == claim_id,
            CaseDocument.source_type == "engineer_detail_upload",
            CaseDocument.is_active == True,  # noqa: E712
            CaseDocument.is_deleted == False,  # noqa: E712
        )
        .order_by(CaseDocument.id.desc())
        .first()
    )
    if report_doc and report_doc.s3_key:
        try:
            report_bytes = s3.read_file_bytes(report_doc.s3_key)
            attachments.append({
                "name": report_doc.original_filename or report_doc.file_name or "Engineer-Report.pdf",
                "content_bytes": base64.b64encode(report_bytes).decode(),
                "content_type": report_doc.content_type or "application/pdf",
            })
        except Exception:
            pass

    # 2) Our covering letter, generated from the claim.
    try:
        cover_bytes = _build_tpi_cover_letter_docx(ctx)
        attachments.append({
            "name": f"Covering Letter - {ctx['reference']}.docx",
            "content_bytes": base64.b64encode(cover_bytes).decode(),
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })
    except Exception:
        pass

    return attachments


def send_engineer_report_to_tpi_email(db: Session, claim_id: int, request, to_email: str = None):
    ctx = _get_claim_context(db, claim_id)

    # Resolve recipient: explicit override → TPI direct_email → error
    tpi_email = to_email or (ctx["tpi"].direct_email if ctx["tpi"] else None)
    receiver_email = _resolve_recipient(tpi_email, None)

    reference = ctx["reference"]
    handler_name = handler_name_for_user(db, actor_id(request)) or handler_name_for_claim(ctx["claim"], db) or (ctx["handler"].label if ctx["handler"] else "Claims Handler")
    ctx["resolved_handler"] = handler_name  # used by the cover-letter helper

    subject = f"Letter Enclosing Copy of Engineers Report for Repair Authority - {reference}"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <body style="margin:0; padding:0; background:#ffffff;">
    <table width="640" cellpadding="0" cellspacing="0" border="0"
           style="margin:0 auto; background:#ffffff; font-family:Arial,sans-serif;">

      <!-- Logo -->
      <tr>
        <td style="padding:40px 30px 24px 30px; text-align:center;">
          <img src="cid:na_logo" width="33" height="31" alt="Logo"
               style="display:inline-block;">
        </td>
      </tr>

      <!-- Dear Sirs -->
      <tr>
        <td style="padding:0 30px 16px 30px; text-align:center;">
          <span style="font-size:16px; font-weight:600; color:#000000;
                       font-family:Arial,sans-serif;">Dear Sirs</span>
        </td>
      </tr>

      <!-- Paragraph 1 -->
      <tr>
        <td style="padding:0 128px 16px 128px; text-align:center;">
          <p style="margin:0; font-size:14px; font-weight:400; color:#444444;
                    line-height:1.6; font-family:Arial,sans-serif;">
            Please find attached the Independent Engineer&apos;s Report
            and our covering letter.
          </p>
        </td>
      </tr>

      <!-- Paragraph 2 -->
      <tr>
        <td style="padding:0 128px 80px 128px; text-align:center;">
          <p style="margin:0; font-size:14px; font-weight:400; color:#444444;
                    line-height:1.6; font-family:Arial,sans-serif;">
            Should you have any questions or queries please do not hesitate
            to contact us on the number provided.
          </p>
        </td>
      </tr>

      <!-- Divider -->
      <tr>
        <td style="padding:0 30px;">
          <div style="height:1px; background:#CCCCCC; width:580px;"></div>
        </td>
      </tr>

      <!-- Footer -->
      <tr>
        <td style="padding:24px 30px 40px 30px; text-align:center;">
          <p style="margin:0; font-size:14px; font-weight:600;
                    color:#444444; font-family:Arial,sans-serif;">
            {handler_name}<br>Claims Handler<br>Nationwide Assist Ltd
          </p>
        </td>
      </tr>
    </table>
    </body>
    </html>
    """

    # The body promises the engineer's report + a covering letter — attach both.
    attachments = _collect_tpi_report_attachments(db, claim_id, ctx)

    status_code = _send_email(receiver_email, subject, html_content, attachments=attachments)

    _log_email_activity(
        db=db,
        request=request,
        claim_id=claim_id,
        file_name=f"Engineer report sent to TPI for claim {reference}",
        file_type=HistoryLogType.ENGINEER_REPORT_SENT_TO_TPI,
    )

    return {"status": "success", "sendgrid_status": status_code}


def send_pav_to_client_email(db: Session, claim_id: int, request, to_email: str = None):
    ctx = _get_claim_context(db, claim_id)
    client_address = ctx["client_address"]
    vehicle = ctx["vehicle"]
    incident = ctx["incident"]

    receiver_email = _resolve_recipient(
        to_email,
        client_address.email if client_address else None,
    )

    reference = ctx["reference"]
    today = datetime.now().strftime("%-d %B %Y")
    handler_name = handler_name_for_user(db, actor_id(request)) or handler_name_for_claim(ctx["claim"], db) or (ctx["handler"].label if ctx["handler"] else "Recovery Handler")
    vehicle_reg = _safe(vehicle.registration if vehicle else None)
    incident_date = _format_date(incident.date_time) if incident else "N/A"

    addr_lines = []
    if client_address:
        if client_address.address:
            addr_lines.append(client_address.address)
        if client_address.postcode:
            addr_lines.append(client_address.postcode)
    addr_block = "<br>".join(addr_lines)

    subject = f"PAV - {reference}"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <body style="margin:0; padding:0; background:#ffffff;">
    <table width="640" cellpadding="0" cellspacing="0" border="0"
           style="margin:0 auto; text-align:center; background:#ffffff; font-family:Arial,sans-serif; color:#444444;">

      <!-- Logo -->
      <tr>
        <td style="padding:40px 30px 30px 30px; text-align:center;">
          <img src="cid:na_logo" width="48" alt="Logo" style="display:inline-block;height:auto;">
        </td>
      </tr>

      <!-- Client address block -->
      <tr>
        <td style="padding:0 30px 20px 30px; font-size:14px; line-height:1.6; text-align:center;">
          <strong>{ctx["salutation"]}</strong><br>
          {addr_block}
        </td>
      </tr>

      <!-- Our ref + date -->
      <tr>
        <td style="padding:0 30px 20px 30px; font-size:14px; line-height:1.6; text-align:center;">
          Our Ref: {_safe(reference)}<br>
          {today}
        </td>
      </tr>

      <!-- Info card -->
      <tr>
        <td style="padding:0 30px 24px 30px; text-align:center;">
          <table width="384" align="center" cellpadding="0" cellspacing="0" border="0"
                 style="border:1px solid #CCCCCC; border-radius:8px; padding:0 16px; background:#ffffff; text-align:left;">
            <tr><td style="padding:8px 0 0 0;"></td></tr>
            <tr>
              <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Incident Date</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{incident_date}</td>
            </tr>
            <tr><td colspan="2" style="height:1px; background:#CCCCCC; padding:0;"></td></tr>
            <tr>
              <td style="width:128px; padding:8px 0; font-size:12px; font-weight:400; color:#444444;">Vehicle</td>
              <td style="padding:8px 0; font-size:12px; font-weight:600; color:#444444;">{vehicle_reg}</td>
            </tr>
            <tr><td style="padding:0 0 8px 0;"></td></tr>
          </table>
        </td>
      </tr>

      <!-- Body -->
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:14px; line-height:1.6;">
          Dear {ctx["salutation"]},
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:14px; line-height:1.6;">
          We are pleased to enclose your cheque in respect of the pre-accident valuation
          of your vehicle involved in the above incident.
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 16px 30px; font-size:14px; line-height:1.6;">
          If you are currently in a credit hire vehicle you may keep the credit hire vehicle
          for a further 7 days from the date we received your cheque.
          Our Fleet Department will be in contact accordingly to arrange the off hire of the vehicle.
        </td>
      </tr>
      <tr>
        <td style="padding:0 30px 40px 30px; font-size:14px; line-height:1.6;">
          If you have any questions or queries, please do not hesitate to contact us
          on the number provided.
        </td>
      </tr>

      <!-- Sign-off -->
      <tr>
        <td style="padding:0 30px 40px 30px; font-size:14px; line-height:1.8;">
          Yours sincerely<br>
          <strong>{handler_name}</strong><br>
          Recovery Handler<br>
          Nationwide Assist Ltd
        </td>
      </tr>

      <!-- Divider -->
      <tr>
        <td style="padding:0 30px;">
          <div style="height:1px; background:#CCCCCC; width:580px;"></div>
        </td>
      </tr>

      <!-- Footer -->
      <tr>
        <td style="padding:24px 30px 40px 30px; text-align:center;">
          <p style="margin:0 0 4px 0; font-size:12px; font-weight:600; color:#444444;">Kind regards,</p>
          <p style="margin:0; font-size:14px; font-weight:600; color:#444444;">
            Nationwide Assist IT / Systems Team
          </p>
        </td>
      </tr>
    </table>
    </body>
    </html>
    """

    status_code = _send_email(receiver_email, subject, html_content)

    _log_email_activity(
        db=db,
        request=request,
        claim_id=claim_id,
        file_name=f"PAV sent to client for claim {reference}",
        file_type=HistoryLogType.PAV_SENT_TO_CLIENT,
    )

    return {"status": "success", "sendgrid_status": status_code}
