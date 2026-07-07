import os
import base64
from io import BytesIO
from datetime import datetime

from fastapi import HTTPException
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail,
    Attachment,
    FileContent,
    FileName,
    FileType,
    Disposition,
    ContentId,
)
from docx import Document

from libdata.enums import PersonRoleEnum, WeatherTypeEnum, HistoryLogType
from libdata.models.tables import (
    ClientDetail,
    VehicleDetail,
    LocationCondition,
    Claim,
    Handler,
    PoliceDetail,
    ClaimQuestionnaire,
    Questionnaire,
    Address,
    ThirdPartyVehicle,
)
from appflow.services.history_activity_service import HistoryActivityService
from appflow.utils import build_case_reference, handler_name_for_claim, handler_name_for_user


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

logo_path = os.path.join(BASE_DIR, "static", "logo.png")
with open(logo_path, "rb") as f:
    logo_encoded = base64.b64encode(f.read()).decode()


def safe(value):
    return str(value) if value is not None else ""


def replace_placeholders_in_docx(doc: Document, data: dict):
    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        changed = False
        for key, value in data.items():
            if key in full_text:
                full_text = full_text.replace(key, safe(value))
                changed = True

        if changed:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def generate_engineer_doc_from_template(data: dict) -> bytes:
    current_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.join(
        current_dir,
        "templates",
        "engineer instruction and letter.docx"
    )

    if not os.path.exists(template_path):
        raise HTTPException(
            status_code=500,
            detail=f"engineer instruction and letter.docx not found at {template_path}"
        )

    doc = Document(template_path)
    replace_placeholders_in_docx(doc, data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()


def send_engineer_instruction_email(
    claim_id: int,
    data,
    db,
    current_user,
    tenant_id
):
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT
    ).first()

    vehicle = db.query(VehicleDetail).filter(
        VehicleDetail.claim_id == claim_id
    ).first()

    accident = db.query(LocationCondition).filter(
        LocationCondition.claim_id == claim_id
    ).first()

    claim = db.query(Claim).filter(
        Claim.id == claim_id
    ).first()

    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    handler = db.query(Handler).filter(
        Handler.id == claim.handler_id
    ).first() if claim.handler_id else None

    police_detail = db.query(PoliceDetail).filter(
        PoliceDetail.claim_id == claim_id
    ).first()

    client_address = db.query(Address).filter(
        Address.id == client.address_id
    ).first() if client and client.address_id else None

    third_vehicle = db.query(ThirdPartyVehicle).filter(
        ThirdPartyVehicle.client_vehicle_id == vehicle.id
    ).first() if vehicle else None

    questionnaire_answer = ""
    claim_questionnaire = db.query(ClaimQuestionnaire).filter(
        ClaimQuestionnaire.claim_id == claim_id
    ).first()

    if claim_questionnaire:
        accident_questionnaire = db.query(Questionnaire).filter(
            Questionnaire.claim_questionnaire_id == claim_questionnaire.id,
            Questionnaire.question.ilike("%accident%") |
            Questionnaire.question.ilike("%description%") |
            Questionnaire.question.ilike("%brief%") |
            Questionnaire.question.ilike("%sketch%")
        ).first()

        if accident_questionnaire:
            questionnaire_answer = accident_questionnaire.answer or ""
        else:
            first_questionnaire = db.query(Questionnaire).filter(
                Questionnaire.claim_questionnaire_id == claim_questionnaire.id
            ).first()

            if first_questionnaire:
                questionnaire_answer = first_questionnaire.answer or ""

    weather_description = ""
    if accident and accident.condition:
        weather_description = (
            accident.condition.value
            if isinstance(accident.condition, WeatherTypeEnum)
            else accident.condition
        )

    year = claim.file_opened_at.strftime("%Y") if claim.file_opened_at else datetime.now().strftime("%Y")
    month = claim.file_opened_at.strftime("%m") if claim.file_opened_at else datetime.now().strftime("%m")
    padded_id = str(claim.id).zfill(5)

    our_reference = (
        f"{client.surname}-{year}{month}-{padded_id}"
        if client and client.surname
        else f"{year}{month}-{padded_id}"
    )

    client_name = f"{client.first_name or ''} {client.surname or ''}".strip() if client else ""

    vehicle_details = (
        f"{vehicle.registration or ''}/{vehicle.make or ''} {vehicle.model or ''}".strip()
        if vehicle else ""
    )

    incident_date = accident.date_time.strftime("%d/%m/%Y") if accident and accident.date_time else ""
    incident_time = accident.date_time.strftime("%H:%M") if accident and accident.date_time else ""

    current_location = data.current_location if data and getattr(data, "current_location", None) else ""

    current_vehicle_status = (
        getattr(vehicle, "vehicle_status", None)
        or getattr(vehicle, "current_status", None)
        or getattr(vehicle, "status", None)
        or ""
    )

    if hasattr(current_vehicle_status, "value"):
        current_vehicle_status = current_vehicle_status.value

    engineer_address_full = "\n".join(
        filter(None, [
            getattr(data, "engineer_company", ""),
            getattr(data, "engineer_address", ""),
            getattr(data, "engineer_postcode", ""),
        ])
    )

    police_details = ""
    if police_detail:
        police_details = " - ".join(
            filter(None, [
                getattr(police_detail, "name", ""),
                getattr(police_detail, "station_name", ""),
            ])
        )

    email_details = {
        "${EngineerAddress}": engineer_address_full,
        "${Date}": datetime.now().strftime("%d/%m/%y"),
        "${OurReference}": our_reference,
        "${ClientName}": client_name,
        "${VehicleDetails}": vehicle_details,
        "${IncidentDate}": incident_date,
        "${CurrentVehicleStatus}": current_vehicle_status,
        "${CurrentLocation}": current_location,
        "${HandlerLabel}": (
            handler_name_for_user(db, current_user)
            or handler_name_for_claim(claim, db)
            or (handler.label if handler else "")
        ),
    }

    report_details = {
        "${EngineerAddress}": engineer_address_full,
        "${Date}": datetime.now().strftime("%d/%m/%y"),
        "${OurReference}": our_reference,
        "${ClientName}": client_name,
        # Signature on the instruction letter = the logged-in claim handler.
        "${HandlerLabel}": email_details["${HandlerLabel}"],

        "${ClientAddress}": client_address.address if client_address else "",
        "${ClientPostCode}": client_address.postcode if client_address else "",
        "${ClientNi}": client.ni_number if client else "",
        "${DOB}": client.date_of_birth.strftime("%d/%m/%Y") if client and client.date_of_birth else "",
        "${HomeTel}": client_address.home_tel if client_address else "",
        "${PhoneNo}": client_address.home_tel if client_address else "",
        "${WorkNo}": getattr(client_address, "work_tel", "") if client_address else "",
        "${MobileTel}": client_address.mobile_tel if client_address else "",
        "${MobileNo}": client_address.mobile_tel if client_address else "",
        "${Email}": client_address.email if client_address else "",

        "${VehicleDetails}": vehicle_details,
        "${InsuranceType}": claim.claim_type.label if claim and claim.claim_type else "",
        "${Make}": vehicle.make if vehicle and vehicle.make else "",
        "${Model}": vehicle.model if vehicle and vehicle.model else "",
        "${Registration}": vehicle.registration if vehicle and vehicle.registration else "",
        "${CurrentVehicleStatus}": current_vehicle_status,
        "${CurrentLocation}": current_location,

        "${IncidentTime}": incident_time,
        "${IncidentDate}": incident_date,
        "${Location}": accident.location if accident and accident.location else "",
        "${PoliceAttend}": "Yes" if accident and getattr(accident, "police_attend", False) else "No",
        "${PoliceDetails}": police_details,
        "${QuestionnaireAnswer}": questionnaire_answer,
        "${WeatherDescription}": weather_description,

        "${ThirdPartyName}": getattr(third_vehicle, "name", "") if third_vehicle else "",
        "${ThirdPartyVehicle}": getattr(third_vehicle, "vehicle", "") if third_vehicle else "",
        "${ThirdPartyRegistration}": getattr(third_vehicle, "registration", "") if third_vehicle else "",
    }

    doc_content = generate_engineer_doc_from_template(report_details)
    encoded_file = base64.b64encode(doc_content).decode()

    subject = f"Instructing Engineer - {our_reference}"

    handler_name = email_details["${HandlerLabel}"] or "Claims Handler"

    # Body mirrors the instruct-engineer template: a short cover note; the full
    # instruction (letter + accident report form) travels as the attachment.
    body = f"""
    <!DOCTYPE html>
    <html>
    <head>
      <meta charset="UTF-8" />
    </head>

    <body style="margin:0; padding:0; background:#ffffff; font-family:Arial, Helvetica, sans-serif;">
      <table width="100%" cellpadding="0" cellspacing="0" border="0" style="background:#ffffff;">
        <tr>
          <td align="center">

            <table width="560" cellpadding="0" cellspacing="0" border="0" style="width:560px; background:#ffffff;">

              <tr>
                <td align="center" style="padding-top:48px; padding-bottom:24px; text-align:center;">
                  <img
                    src="cid:companylogo"
                    alt="Nationwide Assist"
                    width="32"
                    style="display:inline-block; width:32px; height:auto;"
                  />
                </td>
              </tr>

              <tr>
                <td align="center" style="font-size:14px; line-height:22px; color:#2f3a3a; text-align:center;">
                  <div style="font-weight:600; color:#000000;">Dear Sirs</div>

                  <div style="padding-top:18px;">Please find attached our new instruction.</div>

                  <div style="padding-top:18px;">
                    If you are not able to inspect the Client&rsquo;s vehicle within 48 working hours
                    from the date of this instruction please notify us immediately.
                  </div>

                  <div style="padding-top:18px;">
                    If you have any queries please contact us on the number below.
                  </div>
                </td>
              </tr>

              <tr>
                <td style="padding-top:60px;">
                  <table width="100%" cellpadding="0" cellspacing="0" border="0">
                    <tr>
                      <td style="border-top:1px solid #d3d3d3; height:1px; line-height:1px; font-size:1px;">
                        &nbsp;
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>

              <tr>
                <td align="center" style="padding-top:24px; padding-bottom:50px; text-align:center;">
                  <div style="font-size:12px; line-height:18px; font-weight:600; color:#2f3a3a;">
                    {handler_name}<br/>
                    Claims Handler<br/>
                    Nationwide Assist Ltd<br/>
                    T: 0121 766 7515
                  </div>
                </td>
              </tr>

            </table>

          </td>
        </tr>
      </table>
    </body>
    </html>
    """

    # Graph-first so it reaches Outlook (logo auto-attached via cid:companylogo);
    # SendGrid fallback. The engineer document travels as a regular attachment.
    from appflow.services.email_delivery import send_email as deliver_email

    # TEMP: route to the test inbox instead of the real engineer address for now.
    recipient = "marwanationwideassist@outlook.com"  # was: data.engineer_email

    try:
        result = deliver_email(
            to=recipient,
            subject=subject,
            html=body,
            attachments=[{
                "name": "Engineer Instruction & Letter.docx",
                "content_bytes": encoded_file,
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }],
        )

        reference = build_case_reference(claim_id, db)

        HistoryActivityService.create_activity(
            db=db,
            claim_id=claim_id,
            file_name=f"Instruct Engineer Send For Claim {reference}",
            file_path="",
            file_type=HistoryLogType.INSTRUCT_ENGINEER_SEND,
            user_id=current_user,
            tenant_id=tenant_id
        )

        return {
            "status": "success",
            "delivery": result,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))