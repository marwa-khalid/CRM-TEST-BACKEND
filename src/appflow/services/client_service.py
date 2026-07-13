import base64

from fastapi import HTTPException, Request
from sqlalchemy.orm import Session

from appflow.utils import get_tenant_id,actor_id,build_case_reference
from libdata.models.tables import ClientDetail, Address,Claim,LocationCondition,VehicleDetail,Handler,RouteRepair
from appflow.models.client_detail import ClientDetailIn,ClientDisplayLabels,AddressDisplayLabels
from libdata.enums import PersonRoleEnum,HistoryLogType
from datetime import datetime
import os
from io import BytesIO
from docx import Document
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName, FileType, Disposition, ContentId
)
from appflow.services.history_activity_service import HistoryActivityService
from appflow.services.graph_email_service import GraphEmailService

BASE_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
LOGO_PATH = os.path.join(BASE_TEMPLATE_DIR, "logo.png")
with open(LOGO_PATH, "rb") as logo_file:
    LOGO_ENCODED = base64.b64encode(logo_file.read()).decode()
def create_client_service(request: Request, client: ClientDetailIn, db: Session, role: PersonRoleEnum):
    tenant_id = get_tenant_id(request)
    current_user = actor_id(request)
    claim = db.query(Claim).filter(
        Claim.id == client.claim_id,
        Claim.tenant_id == tenant_id
    ).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")
    # Extra validation for CLIENT role
    # if role == PersonRoleEnum.CLIENT:
    #     required_fields = {
    #         "date_of_birth": client.date_of_birth,
    #         "ni_number": client.ni_number,
    #         "sort_code": client.sort_code,
    #         "account_number": client.account_number,
    #         "surname": client.surname,
    #         "language_id": client.language_id,
    #     }
    #     missing = [field for field, value in required_fields.items() if not value]
    #     if missing:
    #         raise HTTPException(
    #             status_code=422,
    #             detail=f"Missing required fields for CLIENT role: {', '.join(missing)}"
    #         )

    # 1. Save Address if provided
    address_id = None
    if client.address:
        db_address = Address(**client.address.dict())
        db.add(db_address)
        db.flush()
        db.refresh(db_address)
        address_id = db_address.id

    # 2. Save ClientDetail
    db_client = ClientDetail(
        **client.dict(exclude={"address", "tenant_id","language_id"}),
        tenant_id=tenant_id,
        address_id=address_id,
        role=role.value if isinstance(role, PersonRoleEnum) else role,
        language_id=client.language_id,
        created_by = current_user,
        updated_by = current_user,
    )

    db.add(db_client)
    db.commit()
    db.refresh(db_client)
    # reference = build_case_reference(claim.id,db)
    current_yyyymm = datetime.now().strftime("%Y%m")
    padded_claim_id = str(claim.id).zfill(5)
    HistoryActivityService.create_activity(
        db=db,
        claim_id=db_client.claim_id,
        file_name=f"The client detail has been created for claim-{current_yyyymm}-{padded_claim_id}",
        file_path="",
        file_type=HistoryLogType.CREATED_CLIENT_DETAIL,
        user_id=current_user,
        tenant_id=tenant_id
    )

    return db_client

def list_clients_service(request: Request, db: Session, role: PersonRoleEnum | None = None):
    tenant_id = get_tenant_id(request)

    query = db.query(ClientDetail).filter(ClientDetail.tenant_id == tenant_id)

    if role:
        query = query.filter(ClientDetail.role == role)

    return query.all()


def get_client_service(client_id: int, request: Request, db: Session, role: PersonRoleEnum | None = None):
    tenant_id = get_tenant_id(request)

    query = db.query(ClientDetail).filter(ClientDetail.id == client_id, ClientDetail.tenant_id == tenant_id)

    if role:
        query = query.filter(ClientDetail.role == role)

    client = query.first()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    return client
def get_client_by_claim_id(claim_id: int, request: Request, db: Session, role: PersonRoleEnum | None = None):
    tenant_id = get_tenant_id(request)

    query = db.query(ClientDetail).filter(ClientDetail.claim_id == claim_id, ClientDetail.tenant_id == tenant_id)

    if role:
        query = query.filter(ClientDetail.role == role)

    client = query.first()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    return client

def update_client_service(claim_id: int, request: Request, client_data: ClientDetailIn, db: Session, role: PersonRoleEnum):
    tenant_id = get_tenant_id(request)
    current_user = actor_id(request)
    db_client = (
        db.query(ClientDetail)
        .filter(ClientDetail.claim_id == claim_id, ClientDetail.tenant_id == tenant_id,ClientDetail.role == PersonRoleEnum.CLIENT)
        .first()
    )

    if not db_client:
        raise HTTPException(status_code=404, detail="Client not found")
    claim = db.query(Claim).filter(
        Claim.id == client_data.claim_id,
        Claim.tenant_id == tenant_id
    ).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")
    # Validation for CLIENT role
    # if role == PersonRoleEnum.CLIENT:
    #     required_fields = {
    #         "date_of_birth": client_data.date_of_birth,
    #         "ni_number": client_data.ni_number,
    #         "sort_code": client_data.sort_code,
    #         "account_number": client_data.account_number,
    #         "surname": client_data.surname,
    #         "language_id": client_data.language_id,
    #     }
    #     missing = [field for field, value in required_fields.items() if not value]
    #     if missing:
    #         raise HTTPException(
    #             status_code=422,
    #             detail=f"Missing required fields for CLIENT role: {', '.join(missing)}"
    #         )

    changed_fields: list[str] = []
    # Update client fields (exclude nested address + relationship fields)
    primitive_fields = client_data.dict(
        exclude={"address", "language", "created_at", "tenant_id"},
        exclude_unset=True
    )

    for key, value in primitive_fields.items():
        if key == "updated_at":
            continue

        old_value = getattr(db_client, key)
        if old_value != value:
            setattr(db_client, key, value)
            changed_fields.append(ClientDisplayLabels.format(key))

    # -----------------------
    # Language
    # -----------------------
    if client_data.language_id and db_client.language_id != client_data.language_id:
        db_client.language_id = client_data.language_id
        changed_fields.append(ClientDisplayLabels.format("language_id"))

    # -----------------------
    # Address
    # -----------------------
    if client_data.address:
        if db_client.address:
            for key, value in client_data.address.dict(exclude_unset=True).items():
                if key in AddressDisplayLabels.EXCLUDE_FIELDS:
                    continue

                old_value = getattr(db_client.address, key)
                if old_value != value:
                    setattr(db_client.address, key, value)
                    changed_fields.append(
                        f"{AddressDisplayLabels.format(key)}"
                    )
        else:
            db_client.address = Address(**client_data.address.dict())
            changed_fields.append("Address Added")

    db_client.updated_at = datetime.now()
    db_client.updated_by = current_user
    db.commit()
    db.refresh(db_client)
    # reference = build_case_reference(claim.id, db)
    current_yyyymm = datetime.now().strftime("%Y%m")
    padded_claim_id = str(claim.id).zfill(5)
    if changed_fields:
        HistoryActivityService.create_activity(
            db=db,
            claim_id=claim_id,
            file_name=f"The client detail has been updated for claim-{current_yyyymm}-{padded_claim_id}",
            file_path=", ".join(changed_fields),
            file_type=HistoryLogType.UPDATED_CLIENT_DETAIL,
            user_id=current_user,
            tenant_id=tenant_id
        )
    return db_client


def deactivate_client_service(claim_id: int, request: Request, db: Session, role: PersonRoleEnum | None = None):
    tenant_id = get_tenant_id(request)

    query = db.query(ClientDetail).filter(ClientDetail.claim_id == claim_id, ClientDetail.tenant_id == tenant_id)

    if role:
        query = query.filter(ClientDetail.role == role)

    db_client = query.first()
    if not db_client:
        raise HTTPException(status_code=404, detail="Client not found")

    db_client.is_active = False
    db.commit()
    db.refresh(db_client)

    return {"detail": "Client deactivated successfully"}


def generate_cil_agreement_letter_from_template(template_path: str, data: dict) -> bytes:
    """
    Load the template docx file and replace placeholders with data.
    """
    doc = Document(template_path)

    def replace_text_in_paragraphs(paragraphs, data):
        for p in paragraphs:
            for key, value in data.items():
                if key in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            inline[i].text = inline[i].text.replace(key, str(value) if value is not None else "")

    replace_text_in_paragraphs(doc.paragraphs, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def prepare_cil_agreement_letter(claim_id: int, db: Session) -> bytes:
    """
    Service function to generate engineer instruction DOCX for a given claim.
    """
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT
    ).first()
    claim = db.query(Claim).filter(Claim.id == claim_id).first()
    client_address = db.query(Address).filter(Address.id == client.address_id).first() if client else None

    route_repair = db.query(RouteRepair).filter(RouteRepair.claim_id == claim_id).first()
    damage_amount = route_repair.sub_total if route_repair and route_repair.sub_total else 0
    # if not client or not claim or not client_address:
    #     raise HTTPException(status_code=404, detail="Required claim details not found")

    year = claim.file_opened_at.strftime("%Y") if claim.file_opened_at else datetime.now().strftime("%Y")
    month = claim.file_opened_at.strftime("%m") if claim.file_opened_at else datetime.now().strftime("%m")
    padded_id = str(claim.id).zfill(5)
    our_reference = f"{client.surname}-{year}{month}-{padded_id}" if client else f"{year}{month}-{padded_id}"

    # Prepare template placeholders
    details = {
        "${Date}": datetime.now().strftime("%d-%m-%y"),
        "${OurReference}": our_reference,
        "${ClientName}": f"{client.first_name} {client.surname}" if client else "",
        "${ClientAddress}": client_address.address if client_address else "",
        "${ClientPostCode}": client_address.postcode if client_address else "",
        "${DamageAmount}": f"{damage_amount:.2f}",
    }

    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, "templates", "cil_agreement_letter.docx")
    return generate_cil_agreement_letter_from_template(template_path, details)

def prepare_send_cil_to_client(claim_id: int, db: Session) -> bytes:
    """
    Service function to generate engineer instruction DOCX for a given claim.
    """
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT
    ).first()
    claim = db.query(Claim).filter(Claim.id == claim_id).first()
    vehicle = db.query(VehicleDetail).filter(VehicleDetail.claim_id == claim_id).first()
    handler = db.query(Handler).filter(Handler.id == claim.handler_id).first() if claim else None
    accident = db.query(LocationCondition).filter(LocationCondition.claim_id == claim_id).first()
    client_address = db.query(Address).filter(Address.id == client.address_id).first() if client else None

    # if not client or not accident or not claim or not handler or not vehicle or not client_address:
    #     raise HTTPException(status_code=404, detail="Required claim details not found")

    year = claim.file_opened_at.strftime("%Y") if claim.file_opened_at else datetime.now().strftime("%Y")
    month = claim.file_opened_at.strftime("%m") if claim.file_opened_at else datetime.now().strftime("%m")
    padded_id = str(claim.id).zfill(5)
    our_reference = f"{client.surname}-{year}{month}-{padded_id}" if client else f"{year}{month}-{padded_id}"

    # Prepare template placeholders
    details = {
        "${Date}": datetime.now().strftime("%d-%m-%y"),
        "${OurReference}": our_reference,
        "${ClientName}": f"{client.first_name} {client.surname}" if client else "",
        "${ClientAddress}": client_address.address if client_address else "",
        "${ClientPostCode}": client_address.postcode if client_address else "",
        "${IncidentDate}": accident.date_time.strftime('%d-%m-%Y') if accident and accident.date_time else "",
        "${Registration}": vehicle.registration if vehicle and vehicle.registration else "",
        "${HandlerLabel}": handler.label if handler else "",
    }

    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, "templates", "send_cil_to_client.docx")
    return generate_cil_agreement_letter_from_template(template_path, details)

def get_static_doc(claim_id: int, filename: str) -> str:
    """
    Service function to return the absolute path of a static DOCX file.
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, "templates", filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"File '{filename}' not found")

    return file_path

def send_vulnerable_notify_manager(claim_id: int, db,current_user:int,tenant_id: int):
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT
    ).first()
    claim = db.query(Claim).filter(Claim.id == claim_id).first()

    # if not client:
    #     raise HTTPException(status_code=404, detail=f"No client found for claim ID {claim_id}")


    receiver_email = "usaleem651@gmail.com"
    client_name = f"{client.first_name} {client.surname}".strip() or "Client" if client else ""
    year = claim.file_opened_at.strftime("%Y") if claim.file_opened_at else datetime.now().strftime("%Y")
    month = claim.file_opened_at.strftime("%m") if claim.file_opened_at else datetime.now().strftime("%m")
    padded_id = str(claim.id).zfill(5)
    our_reference = f"{client.surname}-{year}{month}-{padded_id}" if client else f"{year}{month}-{padded_id}"
    Date = datetime.now().strftime("%d-%m-%y")

    # 2. Email HTML content (editable)
    html_content = f"""
                        <html>
                          <head>
                            <style>
                              body {{
                                font-family: 'Inter';
                                margin: 0;
                                padding: 0;
                              }}
                              .header, .footer {{
                                display: flex;
                                align-items: center;
                                padding: 10px 20px;
                              }}
                              .header {{
                                border-bottom: 1px solid #ddd;
                              }}
                              .footer {{
                                border-top: 1px solid #ddd;
                                font-size: 12px;
                                color: #666;
                                justify-content: center;
                              }}
                              .header img {{
                                width: 40px;
                                height: auto;
                                margin-right: 10px;
                              }}
                              .header p {{
                                font-size: 20px;
                                font-weight: bold;
                                margin: 0;
                              }}
                            </style>
                          </head>
                          <body>
                            <div class="header">
                              <img src="cid:companylogo" alt="ProClaim" />
                              <p>ProClaim</p>
                            </div>
                            <div style="padding: 20px;">
                              <p><strong>Case No</strong> {our_reference}</p>
                              <p><strong>Client Name</strong> {client_name}.</p>
                              <p><strong>Date</strong> {Date}.</p>
                              <p>Please note the above client has been identified as a <strong>vulnerable person</strong>.
                                 Kindly review the case and advise if any additional actions or support are required.</p>
                            <div class="footer">
                              <p>This email was sent to claim@nationwideassist.co.uk. If you'd rather not receive this kind of email, you can unsubscribe or manage your email preferences.<br> <img src="cid:companylogo" alt="ProClaim" width="25" style="vertical-align: middle; margin-right: 6px;" /> <span style="font-size:13px; font-weight:bold; vertical-align: middle;">ProClaim</span> 
                            </div>
                          </body>
                        </html>
                        """
    # 3. Prepare and send
    subject = f"Notification of Vulnerable Person - {our_reference}"

    # Prefer Microsoft Graph (delivered from a real Outlook mailbox, logo inline);
    # fall back to SendGrid only if Graph is unavailable or fails.
    sendgrid_status = None
    if GraphEmailService.is_configured():
        result = GraphEmailService.send_mail(receiver_email, subject, html_content)
        if result is not None:
            sendgrid_status = result.status_code

    if sendgrid_status is None:
        message = Mail(
            from_email="no-replynationwideassist@outlook.com",
            to_emails=receiver_email,
            subject=subject,
            html_content=html_content,
        )
        message.add_attachment(Attachment(
            FileContent(LOGO_ENCODED),
            FileName("logo.png"),
            FileType("image/png"),
            Disposition("inline"),
            ContentId("companylogo"),
        ))
        api_key = os.getenv("SENDGRID_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="SendGrid API key not found in environment")
        sg = SendGridAPIClient(api_key)
        response = sg.send(message)
        sendgrid_status = response.status_code
    # reference = build_case_reference(claim_id,db)
    current_yyyymm = datetime.now().strftime("%Y%m")
    padded_claim_id = str(claim.id).zfill(5)
    HistoryActivityService.create_activity(
        db=db,
        claim_id=claim_id,
        file_name=f"Vulnerable notify manager email sent for claim-{current_yyyymm}-{padded_claim_id}",
        file_path="",
        file_type=HistoryLogType.VULNERABLE_NOTIFY_SEND,
        user_id=current_user,
        tenant_id=tenant_id
    )

    return {
        "status": "success",
        "message": f"Email sent successfully to {receiver_email}",
        "client_name": client_name,
        "claim_id": claim_id,
        "sendgrid_status": sendgrid_status,
        "timestamp": datetime.now().isoformat(),
    }

def build_vulnerable_email_data(claim_id: int, db):
    client = db.query(ClientDetail).filter(
        ClientDetail.claim_id == claim_id,
        ClientDetail.role == PersonRoleEnum.CLIENT
    ).first()

    claim = db.query(Claim).filter(Claim.id == claim_id).first()
    from appflow.utils import handler_name_for_claim
    handler_name = handler_name_for_claim(claim, db)
    if not handler_name and claim and claim.handler:
        handler_name = claim.handler.label or ""

    # ---- SAFE FIELD EXTRACTIONS ----
    client_first = client.first_name or "" if client else ""
    client_last = client.surname or "" if client else ""
    client_name = f"{client_first} {client_last}".strip()

    # If claim not found, fallback to empty strings
    opened_at = claim.file_opened_at if claim and claim.file_opened_at else None

    year = opened_at.strftime("%Y") if opened_at else ""
    month = opened_at.strftime("%m") if opened_at else ""
    padded_id = str(claim.id).zfill(5) if claim and claim.id else ""

    reference = (
        f"{client_last}-{year}{month}-{padded_id}" if client_last else f"{year}{month}-{padded_id}"
    )

    date = datetime.now().strftime("%d-%m-%y")
    subject = f"Notification of Vulnerable Person -{reference}" if reference else ""    # NEW TEXT CONTENT YOU ASKED FOR
    message_body = (
        "Please note the above client has been identified as a vulnerable person. "
        "Kindly review the case and advise if any additional actions or support are required."
    )

    footer_text = (
        "This email was sent to claim@nationwideassist.co.uk. "
        "If you'd rather not receive this kind of email, you can unsubscribe or manage your email preferences."
    )
    to_recipients = "ayesha.rana@nationwideassist.co.uk;ayesha.rana@nationwideassist.co.uk"
    return {
        "to": to_recipients,
        "subject": subject or "",
        "client_name": client_name or "",
        "reference": reference or "",
        "date": date or "",
        "claim_id": claim_id,
        "message_body": message_body or "",
        "footer_text": footer_text or "",
        "handler": handler_name
    }
