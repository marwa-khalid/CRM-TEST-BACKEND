"""Case History service — chronological log of user-recorded case activities
(letters, emails, calls, notes, diary). Backs the History screen: list (with
search + filters), record detail, and create. Kept separate from the file-based
Case Activity (HistoryActivities)."""
import os
import uuid
from datetime import datetime, timedelta
from typing import List, Optional

from sqlalchemy import and_, or_
from sqlalchemy.orm import Session

from libdata.enums import CaseHistoryActionType, PersonRoleEnum
from libdata.models.tables import (
    CaseHistory,
    Claim,
    ThirdPartyInsurer,
    ClientDetail,
    Address,
    User,
    HireVehicleProvided,
)
from appflow.services.microsoft_graph_token_service import MicrosoftGraphTokenService
from appflow.services.outlook_case_activity_service import OutlookCaseActivityService
from appflow.services.s3_service import S3Service
from appflow.services.case_email_import import parse_email_bytes
from appflow.utils import build_case_reference, handler_name_for_user


def _fmt_size(num: int) -> str:
    size = float(num or 0)
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024 or unit == "GB":
            return f"{size:.0f} {unit}" if unit == "B" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} GB"

# wire value ("send_letter") -> enum member
_ACTION_BY_VALUE = {a.value: a for a in CaseHistoryActionType}


def _to_dict(r: CaseHistory) -> dict:
    payload = r.payload
    # Give stored (imported-email) attachments an openable URL routed through the
    # record — the S3 key stays server-side.
    if isinstance(payload, dict) and isinstance(payload.get("attachments"), list):
        atts = []
        for i, a in enumerate(payload["attachments"]):
            a = dict(a or {})
            if a.get("s3_key") and not a.get("url"):
                a["url"] = f"/case-history/{r.id}/attachment/{i}"
            atts.append(a)
        payload = {**payload, "attachments": atts}
    return {
        "id": r.id,
        "claim_id": r.claim_id,
        "scope_type": r.scope_type,
        "scope_id": r.scope_id,
        "action_type": r.action_type.value if r.action_type else None,
        "posted_at": r.posted_at,
        "correspondent": r.correspondent,
        "handler": r.handler,
        "subject": r.subject,
        "details": r.details,
        "payload": payload,
        "created_by": r.created_by,
        "created_at": r.created_at,
    }


def _scope_clause(scope_type: str, scope_id: int):
    """SQLAlchemy filter selecting one owner's records. Claim records also carry
    claim_id, but every row now has scope_type/scope_id, so this works for both."""
    return and_(CaseHistory.scope_type == scope_type, CaseHistory.scope_id == scope_id)


class CaseHistoryService:
    @staticmethod
    def _list(
        db: Session,
        scope_type: str,
        scope_id: int,
        *,
        search: Optional[str] = None,
        action_type: Optional[List[str]] = None,
        correspondent: Optional[List[str]] = None,
        handler: Optional[List[str]] = None,
        date_from: Optional[datetime] = None,
        date_to: Optional[datetime] = None,
    ) -> List[dict]:
        q = db.query(CaseHistory).filter(
            _scope_clause(scope_type, scope_id),
            CaseHistory.is_deleted.isnot(True),
        )
        if action_type:
            members = [_ACTION_BY_VALUE[a] for a in action_type if a in _ACTION_BY_VALUE]
            if members:
                q = q.filter(CaseHistory.action_type.in_(members))
        if correspondent:
            q = q.filter(CaseHistory.correspondent.in_(correspondent))
        if handler:
            q = q.filter(CaseHistory.handler.in_(handler))
        if date_from:
            q = q.filter(CaseHistory.posted_at >= date_from)
        if date_to:
            # `date_to` is an inclusive end-of-day boundary.
            q = q.filter(CaseHistory.posted_at < date_to + timedelta(days=1))
        if search and search.strip():
            like = f"%{search.strip()}%"
            q = q.filter(
                or_(
                    CaseHistory.details.ilike(like),
                    CaseHistory.subject.ilike(like),
                    CaseHistory.correspondent.ilike(like),
                    CaseHistory.handler.ilike(like),
                )
            )
        # Newest first (chronological, most recent at the top).
        q = q.order_by(CaseHistory.posted_at.desc().nullslast(), CaseHistory.id.desc())
        return [_to_dict(r) for r in q.all()]

    @staticmethod
    def list_for_claim(db: Session, claim_id: int, **kw) -> List[dict]:
        return CaseHistoryService._list(db, "claim", claim_id, **kw)

    @staticmethod
    def list_for_scope(db: Session, scope_type: str, scope_id: int, **kw) -> List[dict]:
        return CaseHistoryService._list(db, scope_type, scope_id, **kw)

    @staticmethod
    def get_by_id(db: Session, record_id: int) -> Optional[dict]:
        r = (
            db.query(CaseHistory)
            .filter(CaseHistory.id == record_id, CaseHistory.is_deleted.isnot(True))
            .first()
        )
        return _to_dict(r) if r else None

    @staticmethod
    def create(db: Session, claim_id: int, data, current_user: Optional[int]) -> dict:
        return CaseHistoryService.create_for_scope(
            db, "claim", claim_id, data, current_user,
            tenant_id=db.query(Claim.tenant_id).filter(Claim.id == claim_id).scalar(),
        )

    @staticmethod
    def create_for_scope(
        db: Session,
        scope_type: str,
        scope_id: int,
        data,
        current_user: Optional[int],
        *,
        tenant_id: Optional[int] = None,
    ) -> dict:
        action = _ACTION_BY_VALUE.get((data.action_type or "").strip().lower())
        if action is None:
            raise ValueError(f"Invalid action_type: {data.action_type!r}")
        rec = CaseHistory(
            claim_id=scope_id if scope_type == "claim" else None,
            scope_type=scope_type,
            scope_id=scope_id,
            tenant_id=tenant_id,
            action_type=action,
            posted_at=data.posted_at or datetime.utcnow(),
            correspondent=data.correspondent,
            handler=data.handler,
            subject=data.subject,
            details=data.details,
            payload=data.payload,
            created_by=current_user,
        )
        db.add(rec)
        db.commit()
        db.refresh(rec)
        # Notify any @-mentioned users in the note (same as Case Activity tagging).
        if rec.details and "@" in rec.details:
            try:
                from appflow.services.notification_service import create_mention_notifications
                ref = ""
                try:
                    ref = CaseHistoryService.scope_reference(db, scope_type, scope_id) or ""
                except Exception:
                    ref = ""
                create_mention_notifications(
                    db, note_text=rec.details,
                    claim_id=scope_id if scope_type == "claim" else None,
                    actor_user_id=current_user, tenant_id=tenant_id, case_reference=ref,
                )
            except Exception as exc:  # noqa: BLE001 — never break the history save
                print(f"[CaseHistory] mention notify failed: {exc}")
        # A Diary (DY) entry is pushed to the assigned handler's Task Management queue.
        if action == CaseHistoryActionType.DIARY:
            CaseHistoryService._diary_to_task(db, rec, scope_type, scope_id, tenant_id, current_user)
        # Send Email / Send Letter with the Accident Report Form template → attach the
        # pre-filled PDF.
        if (action in (CaseHistoryActionType.SEND_EMAIL, CaseHistoryActionType.SEND_LETTER)
                and scope_type == "claim"
                and isinstance(rec.payload, dict)
                and (rec.payload.get("template") or "").strip().lower() == "accident report form"):
            CaseHistoryService._attach_accident_report(db, rec, scope_id)
        return _to_dict(rec)

    @staticmethod
    def _diary_to_task(db, rec, scope_type: str, scope_id: int, tenant_id, current_user) -> None:
        """Create a Task Management task from a Diary history entry, assigned to the
        record's handler and due on the diary date, so it lands in that user's queue."""
        try:
            from appflow.models.task import TaskCreate
            from appflow.services.task_service import TaskService

            module = {"claim": "claims", "fleet_hire": "skyline",
                      "vm_cams": "vehicles", "vm_skyline": "vehicles"}.get(scope_type, "")
            pd = rec.payload if isinstance(rec.payload, dict) else {}
            # The claims diary form stores the assignee + due date in the payload;
            # fall back to the record's handler / posted date otherwise.
            assigned = (pd.get("assigned_to") or rec.handler) or None
            title = ((rec.subject or "").strip() or (pd.get("action") or "").strip()
                     or (rec.details or "").strip()[:80] or "Diary reminder")
            due = None
            if pd.get("due_date"):
                try:
                    due = datetime.strptime(str(pd["due_date"])[:10], "%Y-%m-%d").date()
                except Exception:
                    due = None
            if due is None:
                due = rec.posted_at.date() if rec.posted_at else None
            due_time = pd.get("due_time") or None
            ref, reg = None, None
            try:
                ref = CaseHistoryService.scope_reference(db, scope_type, scope_id) or None
            except Exception:
                ref = None
            if scope_type in ("vm_cams", "vm_skyline"):
                reg, ref = ref, None
            payload = TaskCreate(
                title=title,
                description=(rec.details or None),
                assigned_user=assigned,
                module=module,
                due_date=due,
                due_time=due_time,
                priority="Medium",
                status="Pending",
                claim_id=scope_id if scope_type == "claim" else None,
                claim_reference=ref,
                vehicle_registration=reg,
                notes=(rec.details or None),
            )
            TaskService.create_task(payload, db, current_user, tenant_id)
        except Exception as exc:  # noqa: BLE001 — never break the history save
            print(f"[CaseHistory] diary -> task failed: {exc}")

    @staticmethod
    def build_accident_report(db: Session, claim_id: int) -> Optional[bytes]:
        """Fill the Accident Report Form template with the claim's data. Only the
        existing PDF form-field VALUES are set — the template's layout, fonts and
        design are left exactly as-is (nothing about the format changes)."""
        try:
            import io
            import os
            from pypdf import PdfReader, PdfWriter
            from libdata.models.tables import ClientDetail, VehicleDetail, Address, LocationCondition
            from appflow.utils import build_case_reference

            tmpl = os.path.join(os.path.dirname(__file__), "..", "assets", "templates", "AccidentReportForm.pdf")
            if not os.path.exists(tmpl):
                print("[CaseHistory] accident report template missing")
                return None

            claim = db.query(Claim).filter(Claim.id == claim_id).first()
            if not claim:
                return None
            client = (db.query(ClientDetail).filter(ClientDetail.claim_id == claim_id)
                      .order_by(ClientDetail.id).first())
            addr = None
            if client and getattr(client, "address_id", None):
                addr = db.query(Address).filter(Address.id == client.address_id).first()
            vehicle = (db.query(VehicleDetail).filter(VehicleDetail.claim_id == claim_id)
                       .order_by(VehicleDetail.id).first())
            incident = (db.query(LocationCondition).filter(LocationCondition.claim_id == claim_id)
                        .order_by(LocationCondition.id).first())

            def g(o, a):
                return getattr(o, a, None) if o else None

            def dmy(d):
                return d.strftime("%d/%m/%Y") if d else ""

            name = " ".join(x for x in [g(client, "first_name"), g(client, "surname")] if x).strip()
            make_model = " ".join(x for x in [g(vehicle, "make"), g(vehicle, "model")] if x).strip()
            inc_date = g(incident, "date_time")
            try:
                ref = build_case_reference(claim_id, db)
            except Exception:
                ref = ""

            # Field names are generic (untitled1..N); mapping is by position in the form.
            mapped = {
                "untitled1": ref,                                   # Claim Number
                "untitled2": name,                                  # Full Name
                "untitled3": g(client, "occupation") or "",         # Occupation
                "untitled4": dmy(g(client, "date_of_birth")),       # Date of birth
                "untitled5": g(addr, "home_tel") or "",             # Telephone (Home)
                "untitled6": g(addr, "address") or "",              # Address
                "untitled8": g(addr, "postcode") or "",             # Postcode
                "untitled9": g(addr, "landline_tel") or "",         # (Business)
                "untitled10": g(addr, "mobile_tel") or "",          # (Mobile)
                "untitled21": make_model,                           # Make and Model
                "untitled22": g(vehicle, "registration") or "",     # Reg No
                "untitled24": g(vehicle, "color") or "",            # Colour
                "untitled25": g(vehicle, "body_type") or "",        # Type of Body
                "untitled26": g(vehicle, "engine_size") or "",      # Cubic Capacity
                "untitled42": dmy(inc_date.date() if inc_date else None),  # Accident Date
                "untitled44": g(incident, "location") or "",        # Accident Location
            }

            reader = PdfReader(tmpl)
            # The supplied template ships with a completed EXAMPLE (another person's
            # details). Clear every field first so none of that leaks onto this case's
            # report, then fill the fields we have data for. Values only — not layout.
            fields = {k: "" for k in (reader.get_fields() or {}).keys()}
            for k, v in mapped.items():
                if v:
                    fields[k] = v

            writer = PdfWriter()
            writer.append(reader)
            for page in writer.pages:
                try:
                    writer.update_page_form_field_values(page, fields, auto_regenerate=False)
                except Exception:
                    pass
            try:
                writer.set_need_appearances_writer(True)  # make viewers render filled values
            except Exception:
                pass
            buf = io.BytesIO()
            writer.write(buf)
            return buf.getvalue()
        except Exception as exc:  # noqa: BLE001
            print(f"[CaseHistory] accident report fill failed: {exc}")
            return None

    @staticmethod
    def _attach_accident_report(db: Session, rec, claim_id: int) -> None:
        """Generate the pre-filled Accident Report Form and attach it to a send-email
        history record (stored in S3, surfaced as a previewable attachment)."""
        try:
            import uuid
            pdf = CaseHistoryService.build_accident_report(db, claim_id)
            if not pdf:
                return
            fname = "Accident Report Form.pdf"
            att = {"name": fname, "size": _fmt_size(len(pdf))}
            key = f"claims/history/claim/{claim_id}/docs/{uuid.uuid4().hex}_{fname}"
            try:
                s3 = S3Service()
                s3.client.put_object(Bucket=s3.bucket_name, Key=key, Body=pdf,
                                     ContentType="application/pdf")
                att["s3_key"] = key
            except Exception as exc:
                print(f"[CaseHistory] accident report upload failed: {exc}")
            payload = dict(rec.payload or {})
            payload["attachments"] = list(payload.get("attachments") or []) + [att]
            rec.payload = payload
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(rec, "payload")
            db.commit()
        except Exception as exc:  # noqa: BLE001 — never break the history save
            print(f"[CaseHistory] attach accident report failed: {exc}")

    @staticmethod
    def log_document_record(
        db: Session,
        claim_id: int,
        *,
        action_type: CaseHistoryActionType,
        subject: Optional[str] = None,
        details: Optional[str] = None,
        correspondent: Optional[str] = None,
        handler: Optional[str] = None,
        body_html: Optional[str] = None,
        documents: Optional[List[dict]] = None,   # [{name, data: bytes, content_type}]
        message_id: Optional[str] = None,
        source: str = "email",
        current_user: Optional[int] = None,
        posted_at: Optional[datetime] = None,
        scope_type: str = "claim",
        scope_id: Optional[int] = None,
    ) -> dict:
        """Log a Case History record for a system action (engineer instruct email →
        SE, payment-pack download → SL, generated hire document → SL, …). Any
        documents are stored in S3 and surfaced as openable attachments in the detail
        pane. Best-effort upload — the record is still written if S3 fails."""
        if scope_id is None:
            scope_id = claim_id
        is_claim = scope_type == "claim"

        claim_row = db.query(Claim).filter(Claim.id == scope_id).first() if is_claim else None
        tenant_id = claim_row.tenant_id if claim_row else None

        # Claim scope only: auto-fill handler (General Details handler) + correspondent
        # (the claim's third-party email) when the caller didn't supply them.
        if is_claim:
            if handler is None and claim_row is not None:
                h = getattr(claim_row, "handler", None)
                handler = (getattr(h, "label", "") or "") or None
            if correspondent is None:
                correspondent = CaseHistoryService._claim_tpi_correspondent(db, scope_id)
        # Fall back to the user who performed the action when there's no handler on the
        # claim (or for fleet records) — otherwise the row shows a bare "-".
        if handler is None and current_user:
            handler = handler_name_for_user(db, current_user) or None

        att_refs: List[dict] = []
        docs = [d for d in (documents or []) if d.get("data")]
        if docs:
            s3 = S3Service()
            for doc in docs:
                name = doc.get("name") or "document"
                blob = doc.get("data") or b""
                # NB: the S3 IAM policy only grants access under the "claims/" prefix.
                key = f"claims/history/{scope_type}/{scope_id}/docs/{uuid.uuid4().hex}_{name}"
                try:
                    s3.client.put_object(
                        Bucket=s3.bucket_name,
                        Key=key,
                        Body=blob,
                        ContentType=doc.get("content_type") or "application/octet-stream",
                    )
                    att_refs.append({"name": name, "size": _fmt_size(len(blob)), "s3_key": key})
                except Exception as exc:
                    print(f"[CaseHistoryService] document upload failed: {exc}")
                    att_refs.append({"name": name, "size": _fmt_size(len(blob))})

        rec = CaseHistory(
            claim_id=scope_id if is_claim else None,
            scope_type=scope_type,
            scope_id=scope_id,
            tenant_id=tenant_id,
            action_type=action_type,
            posted_at=posted_at or datetime.utcnow(),
            correspondent=correspondent,
            handler=handler,
            subject=subject,
            details=details,
            payload={
                "source": source,
                "message_id": message_id or None,
                "body_html": body_html or "",
                "body_text": "",
                "attachments": att_refs,
            },
            created_by=current_user,
        )
        db.add(rec)
        db.commit()
        db.refresh(rec)
        return _to_dict(rec)

    @staticmethod
    def correspondents(db: Session, claim_id: Optional[int] = None, tenant_id: Optional[int] = None) -> List[dict]:
        """ALL third-party email addresses in the tenant (from every claim's Third
        Party Insurer screen) — each with its name + phone, for the History forms'
        Correspondent dropdown and the outgoing-call phone auto-fill. Scoped to the
        given claim's tenant (or an explicit tenant_id); only entries with an email."""
        if tenant_id is None and claim_id is not None:
            tenant_id = db.query(Claim.tenant_id).filter(Claim.id == claim_id).scalar()

        # Every Third Party Insurer record in the tenant (join Claim for scoping).
        q = (
            db.query(ThirdPartyInsurer)
            .join(Claim, Claim.id == ThirdPartyInsurer.claim_id)
            .filter(ThirdPartyInsurer.is_deleted.isnot(True), Claim.is_deleted.isnot(True))
        )
        if tenant_id is not None:
            q = q.filter(Claim.tenant_id == tenant_id)
        tpis = q.all()

        # Correspondent = the "Email Address" from each case's Third Party Details
        # (ThirdPartyInsurer.third_party_id → Address.email) — one per case. The
        # insurer / handling-agent / direct-email addresses are deliberately excluded.
        client_ids = {tpi.third_party_id for tpi in tpis if tpi.third_party_id}

        out: List[dict] = []
        if client_ids:
            rows = (
                db.query(ClientDetail, Address)
                .outerjoin(Address, Address.id == ClientDetail.address_id)
                .filter(ClientDetail.id.in_(client_ids))
                .all()
            )
            for cd, addr in rows:
                email = (addr.email or "").strip() if addr and addr.email else None
                if not email:
                    continue
                name = " ".join(x for x in [cd.first_name, cd.surname] if x).strip() or None
                phone = (addr.mobile_tel or addr.landline_tel or addr.home_tel or "").strip() or None if addr else None
                role = getattr(cd.role, "value", None) or (cd.role if isinstance(cd.role, str) else None) or "Third Party"
                out.append({"role": role, "name": name, "email": email, "phone": phone})

        # Only correspondents with an email, de-duplicated by email, alphabetical.
        seen, result = set(), []
        for c in sorted(out, key=lambda x: (x["email"] or "").lower()):
            e = c.get("email")
            if e and e.lower() not in seen:
                seen.add(e.lower())
                result.append(c)
        return result

    @staticmethod
    def _cams_vehicle_claim_id(db: Session, vehicle_record_id: int) -> Optional[int]:
        """The claim a CAMS vehicle is on hire against — matched by registration
        number to the claims-side hire vehicle (there's no FK link on the fleet side)."""
        from fleet.models.tables import FleetVehicleRecord
        rec = db.query(FleetVehicleRecord).filter(FleetVehicleRecord.id == vehicle_record_id).first()
        reg = (getattr(rec, "registration_number", "") or "").replace(" ", "").upper()
        if not reg:
            return None
        rows = (
            db.query(HireVehicleProvided.claim_id, HireVehicleProvided.hire_vehicle_registration)
            .filter(HireVehicleProvided.hire_vehicle_registration.isnot(None))
            .all()
        )
        for claim_id, hv_reg in rows:
            if (hv_reg or "").replace(" ", "").upper() == reg:
                return claim_id
        return None

    @staticmethod
    def _claim_client_email(db: Session, claim_id: int) -> Optional[str]:
        """The claim's client (driver) email from Client Details."""
        row = (
            db.query(ClientDetail, Address)
            .outerjoin(Address, Address.id == ClientDetail.address_id)
            .filter(ClientDetail.claim_id == claim_id, ClientDetail.role == PersonRoleEnum.CLIENT)
            .first()
        )
        if row:
            _cd, addr = row
            if addr and (addr.email or "").strip():
                return addr.email.strip()
        return None

    @staticmethod
    def _scope_tenant_id(db: Session, scope_type: str, scope_id: int):
        """Tenant id for a fleet scope, so tenant-wide correspondent options can load."""
        try:
            if scope_type == "fleet_hire":
                from fleet.models.tables import FleetHire
                return db.query(FleetHire.tenant_id).filter(FleetHire.id == scope_id).scalar()
            from fleet.models.tables import FleetVehicleRecord
            return db.query(FleetVehicleRecord.tenant_id).filter(FleetVehicleRecord.id == scope_id).scalar()
        except Exception:
            return None

    @staticmethod
    def _on_hire_drivers(db: Session, tenant_id) -> List[dict]:
        """All drivers currently on hire (hires with an on_hire vehicle) — name + email."""
        try:
            from sqlalchemy import func as _func
            from fleet.models.tables import FleetHire, FleetHireVehicle
            hire_ids = [row[0] for row in db.query(FleetHireVehicle.hire_id)
                        .filter(_func.lower(_func.coalesce(FleetHireVehicle.hire_status, "")) == "on_hire")
                        .distinct().all()]
            if not hire_ids:
                return []
            q = db.query(FleetHire).filter(FleetHire.id.in_(hire_ids), FleetHire.is_deleted.isnot(True))
            if tenant_id is not None:
                q = q.filter(FleetHire.tenant_id == tenant_id)
            out = []
            for h in q.all():
                name = (getattr(h, "driver_name", "") or "").strip()
                email = (getattr(h, "driver_email", "") or "").strip()
                if name or email:
                    out.append({"name": name, "email": email})
            return out
        except Exception as exc:  # noqa: BLE001
            print(f"[CaseHistory] on-hire drivers failed: {exc}")
            return []

    @staticmethod
    def scope_correspondents(db: Session, scope_type: str, scope_id: int) -> dict:
        """Correspondent options for the History correspondent field, per scope.
        Returns {"default": <str|null>, "options": [{"label","value"}]}.
        - claim: tenant-wide Third Party emails.
        - vm_cams: the linked claim's Client email (default) + Third Party emails.
        - vm_skyline / fleet_hire: every driver currently on hire (name shown, email stored).
        """
        default = None
        options: List[dict] = []

        def _email_opts(emails):
            return [{"label": e, "value": e} for e in emails if e]

        if scope_type == "claim":
            options = _email_opts([c["email"] for c in CaseHistoryService.correspondents(db, scope_id) if c.get("email")])
        elif scope_type == "vm_cams":
            claim_id = CaseHistoryService._cams_vehicle_claim_id(db, scope_id)
            if claim_id:
                default = CaseHistoryService._claim_client_email(db, claim_id)
                tps = [c["email"] for c in CaseHistoryService.correspondents(db, claim_id) if c.get("email")]
                options = _email_opts(([default] if default else []) + tps)
            else:
                tid = CaseHistoryService._scope_tenant_id(db, scope_type, scope_id)
                options = _email_opts([c["email"] for c in CaseHistoryService.correspondents(db, tenant_id=tid) if c.get("email")])
        else:  # vm_skyline / fleet_hire — all on-hire drivers' EMAIL addresses.
            tid = CaseHistoryService._scope_tenant_id(db, scope_type, scope_id)
            options = _email_opts([d.get("email") for d in CaseHistoryService._on_hire_drivers(db, tid) if d.get("email")])
        # De-dup by value, preserve order (default/client first).
        seen, out = set(), []
        for o in options:
            k = (o["value"] or "").lower()
            if o["value"] and k not in seen:
                seen.add(k)
                out.append(o)
        return {"default": default, "options": out}

    @staticmethod
    def _claim_tpi_correspondent(db: Session, claim_id: int) -> Optional[str]:
        """The correspondent auto-filled on payment-pack / letter records: the
        "Email Address" from the claim's Third Party Details
        (ThirdPartyInsurer.third_party_id → Address.email)."""
        tpi = (
            db.query(ThirdPartyInsurer)
            .filter(ThirdPartyInsurer.claim_id == claim_id, ThirdPartyInsurer.is_deleted.isnot(True))
            .first()
        )
        if not tpi or not tpi.third_party_id:
            return None
        row = (
            db.query(ClientDetail, Address)
            .outerjoin(Address, Address.id == ClientDetail.address_id)
            .filter(ClientDetail.id == tpi.third_party_id)
            .first()
        )
        if row:
            _cd, addr = row
            if addr and (addr.email or "").strip():
                return addr.email.strip()
        return None

    @staticmethod
    def _other_party(from_email: str, to_list: List[str]) -> Optional[str]:
        """Correspondent = the party that isn't one of our own addresses."""
        ours = {"no-replynationwideassist@outlook.com"}
        mb = (MicrosoftGraphTokenService.mailbox_user("read") or "").strip().lower()
        if mb:
            ours.add(mb)
        ours.add((os.getenv("SENDGRID_SENDER") or "no-replynationwideassist@outlook.com").strip().lower())
        frm = (from_email or "").strip()
        clean_to = [a for a in (to_list or []) if a]
        if frm.lower() in ours:
            for a in clean_to:
                if a.strip().lower() not in ours:
                    return a
            return clean_to[0] if clean_to else (frm or None)
        return frm or (clean_to[0] if clean_to else None)

    @staticmethod
    def import_email(
        db: Session, claim_id: int, filename: str, data: bytes, current_user: Optional[int],
        *, scope_type: str = "claim", scope_id: Optional[int] = None,
    ) -> dict:
        """Parse a dragged-in .eml/.msg, store its attachments in S3, and create an
        Incoming Email (IE) History record against the claim/fleet entity."""
        if scope_id is None:
            scope_id = claim_id
        parsed = parse_email_bytes(filename, data)

        s3 = S3Service()
        att_refs: List[dict] = []
        for att in parsed.get("attachments", []):
            name = att.get("name") or "attachment"
            blob = att.get("data") or b""
            # NB: the S3 IAM policy only grants access under the "claims/" prefix.
            key = f"claims/history/{scope_type}/{scope_id}/emails/{uuid.uuid4().hex}_{name}"
            try:
                s3.client.put_object(
                    Bucket=s3.bucket_name,
                    Key=key,
                    Body=blob,
                    ContentType=att.get("content_type") or "application/octet-stream",
                )
            except Exception as exc:
                print(f"[CaseHistoryService] email attachment upload failed: {exc}")
                continue
            att_refs.append({"name": name, "size": _fmt_size(len(blob)), "s3_key": key})

        from_email = parsed.get("from_email") or ""
        to_list = parsed.get("to") or []
        correspondent = CaseHistoryService._other_party(from_email, to_list)

        handler = None
        if current_user:
            u = db.query(User).filter(User.id == current_user).first()
            if u:
                un = u.user_name or ""
                handler = un.split("@")[0] if "@" in un else (un or None)

        body_text = parsed.get("body_text") or ""
        is_claim = scope_type == "claim"
        tenant_id = (
            db.query(Claim.tenant_id).filter(Claim.id == scope_id).scalar() if is_claim else None
        )
        rec = CaseHistory(
            claim_id=scope_id if is_claim else None,
            scope_type=scope_type,
            scope_id=scope_id,
            tenant_id=tenant_id,
            action_type=CaseHistoryActionType.INCOMING_EMAIL,
            posted_at=parsed.get("date") or datetime.utcnow(),
            correspondent=correspondent,
            handler=handler,
            subject=parsed.get("subject") or None,
            details=(body_text[:500] or None),
            payload={
                "source": "imported_email",
                "from_name": parsed.get("from_name"),
                "from_email": from_email,
                "to": to_list,
                "body_text": body_text,
                "body_html": parsed.get("body_html") or "",
                "attachments": att_refs,
            },
            created_by=current_user,
        )
        db.add(rec)
        db.commit()
        db.refresh(rec)
        return _to_dict(rec)

    @staticmethod
    def attachment_bytes(db: Session, record_id: int, index: int):
        """(bytes, filename, content_type) for a stored record's Nth attachment."""
        rec = (
            db.query(CaseHistory)
            .filter(CaseHistory.id == record_id, CaseHistory.is_deleted.isnot(True))
            .first()
        )
        if not rec or not isinstance(rec.payload, dict):
            return None
        atts = rec.payload.get("attachments") or []
        if index < 0 or index >= len(atts):
            return None
        key = (atts[index] or {}).get("s3_key")
        name = (atts[index] or {}).get("name") or "attachment"
        if not key:
            return None
        s3 = S3Service()
        try:
            obj = s3.client.get_object(Bucket=s3.bucket_name, Key=key)
            return (obj["Body"].read(), name, obj.get("ContentType") or "application/octet-stream")
        except Exception as exc:
            print(f"[CaseHistoryService] attachment fetch failed: {exc}")
            return None

    @staticmethod
    def attachment_preview_pages(db: Session, record_id: int, index: int) -> dict:
        """Render a stored attachment as page images (PDF → one PNG per page, via
        PyMuPDF) so the History detail can show the document like the Document
        Library — no browser PDF viewer. Images are returned as data URIs."""
        import base64

        result = CaseHistoryService.attachment_bytes(db, record_id, index)
        if not result:
            return {"type": "unsupported", "pages": []}
        raw, name, content_type = result
        lname = (name or "").lower()
        ctype = (content_type or "").lower()

        if ctype.startswith("image/") or lname.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
            b64 = base64.b64encode(raw).decode("utf-8")
            mime = ctype if ctype.startswith("image/") else "image/png"
            return {"type": "image", "file_name": name, "url": f"data:{mime};base64,{b64}", "pages": []}

        if "pdf" in ctype or lname.endswith(".pdf"):
            pages = CaseHistoryService._pdf_bytes_to_pages(raw)
            if pages is None:
                return {"type": "unsupported", "file_name": name, "pages": []}
            return {"type": "pdf", "file_name": name, "pages": pages}

        # Excel → an exact snapshot of the real sheet (borders, fills, merged cells,
        # images, layout) via headless LibreOffice. SinglePageSheets puts each sheet on
        # ONE page so the sheet isn't chopped across pages; pages render to images and
        # scroll in the preview box. Falls back to the openpyxl HTML grid if LibreOffice
        # isn't available.
        if "spreadsheetml" in ctype or lname.endswith((".xlsx", ".xls")):
            is_legacy_xls = lname.endswith(".xls") and "spreadsheetml" not in ctype
            suffix = ".xls" if is_legacy_xls else ".xlsx"
            pdf_bytes = CaseHistoryService._office_to_pdf(
                raw, suffix,
                convert_to='pdf:calc_pdf_Export:{"SinglePageSheets":{"type":"boolean","value":"true"}}',
            )
            if pdf_bytes:
                pages = CaseHistoryService._pdf_bytes_to_pages(pdf_bytes, crop=True)
                if pages:
                    return {"type": "pdf", "file_name": name, "pages": pages}
            try:
                html = (CaseHistoryService._xls_to_html(raw) if is_legacy_xls
                        else CaseHistoryService._xlsx_to_html(raw))
                return {"type": "html", "file_name": name, "html": html, "pages": []}
            except Exception as exc:
                print(f"[CaseHistoryService] Excel grid preview failed: {exc}")
                return {"type": "unsupported", "file_name": name, "pages": []}

        # Office documents (Word / PowerPoint). Preferred path: convert to PDF with
        # headless LibreOffice and render exact page-image snapshots — pixel-for-pixel
        # the real document, same quality as the PDF preview. When LibreOffice isn't
        # installed we fall back to a best-effort HTML rendering.
        is_word = "wordprocessingml" in ctype or lname.endswith(".docx")
        is_excel = False  # Excel handled above as a scrollable HTML grid
        is_ppt = "presentationml" in ctype or lname.endswith((".pptx", ".ppt"))
        is_legacy_doc = "msword" in ctype or lname.endswith(".doc")

        # The payment-pack "Word" download is actually an HTML document with a .doc
        # extension — render its HTML as-is (already a designed template).
        if is_legacy_doc:
            text = raw.decode("utf-8", "ignore")
            if "<html" in text.lower() or "<body" in text.lower():
                return {"type": "html", "file_name": name, "html": text, "pages": []}

        if is_word or is_excel or is_ppt or is_legacy_doc:
            suffix = "." + (lname.rsplit(".", 1)[-1] if "." in lname else
                            ("docx" if is_word else "xlsx" if is_excel else "pptx"))
            pdf_bytes = CaseHistoryService._office_to_pdf(raw, suffix)
            if pdf_bytes:
                pages = CaseHistoryService._pdf_bytes_to_pages(pdf_bytes, crop=not is_excel)
                if pages:
                    return {"type": "pdf", "file_name": name, "pages": pages}
            # LibreOffice unavailable / failed → HTML fallback.
            try:
                if is_word:
                    return {"type": "html", "file_name": name, "html": CaseHistoryService._docx_to_html(raw), "pages": []}
                if is_excel and lname.endswith(".xlsx"):
                    return {"type": "html", "file_name": name, "html": CaseHistoryService._xlsx_to_html(raw), "pages": []}
            except Exception as exc:
                print(f"[CaseHistoryService] Office HTML fallback failed: {exc}")
            return {"type": "unsupported", "file_name": name, "pages": []}

        return {"type": "unsupported", "file_name": name, "pages": []}

    @staticmethod
    def _soffice_bin() -> Optional[str]:
        """Locate the headless LibreOffice binary, if installed."""
        import os
        import shutil
        for cand in ("soffice", "libreoffice"):
            found = shutil.which(cand)
            if found:
                return found
        for path in (
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS cask
            "/usr/bin/soffice",
            "/usr/lib/libreoffice/program/soffice",  # nixpacks / debian
        ):
            if os.path.exists(path):
                return path
        return None

    @staticmethod
    def _office_to_pdf(raw: bytes, suffix: str, convert_to: str = "pdf") -> Optional[bytes]:
        """Convert an Office document (docx/xlsx/pptx/…) to PDF bytes via headless
        LibreOffice so it can be rendered as exact page-image snapshots. ``convert_to``
        lets the caller pass a filtered target (e.g. the Calc export with
        SinglePageSheets so each sheet lands on ONE page). Returns None when
        LibreOffice isn't installed or the conversion fails (caller falls back)."""
        soffice = CaseHistoryService._soffice_bin()
        if not soffice:
            return None
        import glob
        import os
        import subprocess
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, f"in{suffix}")
            with open(src, "wb") as fh:
                fh.write(raw)
            profile = os.path.join(tmp, "profile")  # isolated per-call user profile
            try:
                subprocess.run(
                    [soffice, "--headless", "--norestore", "--nolockcheck", "--nodefault",
                     f"-env:UserInstallation=file://{profile}",
                     "--convert-to", convert_to, "--outdir", tmp, src],
                    check=True, capture_output=True, timeout=120,
                )
            except Exception as exc:  # noqa: BLE001
                print(f"[CaseHistoryService] LibreOffice convert failed: {exc}")
                return None
            pdfs = glob.glob(os.path.join(tmp, "*.pdf"))
            if not pdfs:
                return None
            with open(pdfs[0], "rb") as fh:
                return fh.read()

    @staticmethod
    def _pdf_bytes_to_pages(raw: bytes, crop: bool = True) -> Optional[list]:
        """Render PDF bytes to a list of {page, image(data-uri PNG)} via PyMuPDF.
        ``crop`` trims surrounding white margins and skips fully-blank pages (good for
        letters/documents); pass crop=False for spreadsheets where the sheet grid /
        page extent should be preserved. Returns None on a render error."""
        import base64
        import io
        try:
            import fitz
            from PIL import Image, ImageChops
        except Exception as exc:  # noqa: BLE001
            print(f"[CaseHistoryService] PDF render deps missing: {exc}")
            return None
        try:
            pdf = fitz.open(stream=raw, filetype="pdf")
            pages = []
            page_no = 0
            for i in range(len(pdf)):
                pix = pdf.load_page(i).get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                if crop:
                    bg = Image.new("RGB", img.size, (255, 255, 255))
                    bbox = ImageChops.difference(img, bg).getbbox()
                    if not bbox:
                        continue  # entirely blank page — skip
                    pad = 12
                    l, t, r2, b2 = bbox
                    img = img.crop((max(0, l - pad), max(0, t - pad),
                                    min(img.width, r2 + pad), min(img.height, b2 + pad)))
                out = io.BytesIO()
                img.save(out, "PNG")
                page_no += 1
                b64 = base64.b64encode(out.getvalue()).decode("utf-8")
                pages.append({"page": page_no, "image": f"data:image/png;base64,{b64}"})
            pdf.close()
            return pages
        except Exception as exc:  # noqa: BLE001
            print(f"[CaseHistoryService] PDF preview render failed: {exc}")
            return None

    @staticmethod
    def _docx_to_html(data: bytes) -> str:
        """Render a .docx as HTML, preserving run formatting (bold/italic/underline),
        paragraph alignment and tables in document order. Not a pixel-perfect Word
        render (that needs LibreOffice) but keeps the document's structure + emphasis."""
        import io
        import html as _html
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(io.BytesIO(data))
        align = {1: "center", 2: "right", 3: "justify"}  # WD_ALIGN_PARAGRAPH values

        def render_para(p) -> str:
            runs = []
            for run in p.runs:
                t = _html.escape(run.text)
                if not t:
                    continue
                if run.bold:
                    t = f"<strong>{t}</strong>"
                if run.italic:
                    t = f"<em>{t}</em>"
                if run.underline:
                    t = f"<u>{t}</u>"
                runs.append(t)
            inner = "".join(runs)
            if not inner.strip():
                return ""  # skip empty paragraphs — they only add white space
            a = align.get(getattr(p, "alignment", None), "left")
            return f"<p style='margin:3px 0;text-align:{a}'>{inner}</p>"

        parts: List[str] = []
        for child in doc.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                html_p = render_para(Paragraph(child, doc))
                if html_p:
                    parts.append(html_p)
            elif tag == "tbl":
                rows = []
                for row in Table(child, doc).rows:
                    cells = "".join(
                        "<td style='border:1px solid #d1d5db;padding:4px 8px;vertical-align:top'>"
                        + (("".join(filter(None, [render_para(cp) for cp in cell.paragraphs]))) or "&nbsp;")
                        + "</td>"
                        for cell in row.cells
                    )
                    rows.append(f"<tr>{cells}</tr>")
                parts.append(f"<table style='border-collapse:collapse;margin:10px 0;width:100%'>{''.join(rows)}</table>")
        return (
            "<div style=\"font-family:Arial,Helvetica,sans-serif;font-size:14px;color:#111827;line-height:1.5\">"
            + "".join(parts) + "</div>"
        )

    @staticmethod
    def _theme_palette(wb) -> list:
        """The 12 theme colours (as #rrggbb), ordered by the SpreadsheetML theme index
        (0=background1, 1=text1, 2=background2, 3=text2, 4-9=accent1-6, 10=hlink,
        11=folHlink). Parsed from the workbook's theme XML; falls back to the Office
        default palette so theme-coloured cells/fonts still resolve when missing."""
        import xml.etree.ElementTree as ET
        # Office default (Excel "Office" theme) — used if the file has no theme XML.
        default = ["FFFFFF", "000000", "E7E6E6", "44546A", "4472C4", "ED7D31",
                   "A5A5A5", "FFC000", "5B9BD5", "70AD47", "0563C1", "954F72"]
        raw = getattr(wb, "loaded_theme", None)
        if not raw:
            return ["#" + c for c in default]
        try:
            ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            root = ET.fromstring(raw)
            scheme = root.find(f".//{ns}clrScheme")
            seq = []  # file order: dk1, lt1, dk2, lt2, accent1-6, hlink, folHlink
            for child in scheme:
                srgb = child.find(f"{ns}srgbClr")
                sysc = child.find(f"{ns}sysClr")
                if srgb is not None:
                    seq.append(srgb.get("val", "000000"))
                elif sysc is not None:
                    seq.append(sysc.get("lastClr", sysc.get("val", "000000")))
                else:
                    seq.append("000000")
            if len(seq) >= 12:
                # Reorder to the theme-index mapping (0/1 and 2/3 are swapped vs. file).
                order = [seq[1], seq[0], seq[3], seq[2], seq[4], seq[5],
                         seq[6], seq[7], seq[8], seq[9], seq[10], seq[11]]
                return ["#" + c.lstrip("#")[-6:].upper() for c in order]
        except Exception:
            pass
        return ["#" + c for c in default]

    @staticmethod
    def _apply_tint(hex_color: str, tint: float) -> str:
        """Lighten (tint>0) / darken (tint<0) a #rrggbb colour the way Excel does."""
        import colorsys
        if not tint:
            return hex_color
        h = hex_color.lstrip("#")
        try:
            r, g, b = int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255
        except Exception:
            return hex_color
        hh, ll, ss = colorsys.rgb_to_hls(r, g, b)
        ll = ll * (1 + tint) if tint < 0 else ll * (1 - tint) + tint
        ll = max(0.0, min(1.0, ll))
        r, g, b = colorsys.hls_to_rgb(hh, ll, ss)
        return "#%02X%02X%02X" % (round(r * 255), round(g * 255), round(b * 255))

    @staticmethod
    def _xlsx_to_html(data: bytes) -> str:
        """Render a .xlsx workbook as an Excel-style HTML grid — column letters
        (A, B, C…) and row numbers, real cell fills (incl. theme / indexed / dark
        colours), per-side borders, font styling, merged cells, column widths and
        embedded pictures — inside a scrollable box so it reads like the real
        spreadsheet and scrolls both ways. Powered by openpyxl."""
        import io
        import base64
        import html as _html
        from datetime import datetime, date, time
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        try:
            from openpyxl.styles.colors import COLOR_INDEX
        except Exception:
            COLOR_INDEX = ()

        wb = load_workbook(io.BytesIO(data), data_only=True)
        theme = CaseHistoryService._theme_palette(wb)
        MAX_ROWS, MAX_COLS = 500, 60
        ROWHDR_W, HROW_PX, DEF_COL_PX, DEF_ROW_PX = 42, 22, 82, 20

        def resolve_color(c):
            """openpyxl Color → #rrggbb, honouring rgb / theme(+tint) / indexed. None
            for auto / unset so callers can skip it."""
            if c is None:
                return None
            try:
                ctype = getattr(c, "type", None)
                if ctype == "rgb" or getattr(c, "rgb", None):
                    rgb = c.rgb
                    if isinstance(rgb, str) and len(rgb) >= 6:
                        if len(rgb) == 8 and rgb[:2].upper() == "00":
                            return None  # fully-transparent
                        base = "#" + rgb[-6:].upper()
                        return CaseHistoryService._apply_tint(base, getattr(c, "tint", 0.0) or 0.0)
                if ctype == "theme":
                    idx = getattr(c, "theme", None)
                    if isinstance(idx, int) and 0 <= idx < len(theme):
                        return CaseHistoryService._apply_tint(theme[idx], getattr(c, "tint", 0.0) or 0.0)
                if ctype == "indexed":
                    idx = getattr(c, "indexed", None)
                    if isinstance(idx, int) and 0 <= idx < len(COLOR_INDEX):
                        argb = COLOR_INDEX[idx]
                        if isinstance(argb, str) and len(argb) >= 6:
                            return "#" + argb[-6:].upper()
            except Exception:
                pass
            return None

        BORDER_STYLE = {
            "thin": ("1px", "solid"), "hair": ("1px", "solid"), "medium": ("2px", "solid"),
            "thick": ("3px", "solid"), "double": ("3px", "double"), "dotted": ("1px", "dotted"),
            "dashed": ("1px", "dashed"), "mediumDashed": ("2px", "dashed"),
            "dashDot": ("1px", "dashed"), "mediumDashDot": ("2px", "dashed"),
            "dashDotDot": ("1px", "dashed"), "mediumDashDotDot": ("2px", "dashed"),
            "slantDashDot": ("2px", "dashed"),
        }

        def side_border(side):
            if side is None or not getattr(side, "style", None):
                return None
            w, s = BORDER_STYLE.get(side.style, ("1px", "solid"))
            col = resolve_color(getattr(side, "color", None)) or "#000000"
            return f"{w} {s} {col}"

        def is_dark(hex_color):
            try:
                h = hex_color.lstrip("#")
                lum = 0.299 * int(h[0:2], 16) + 0.587 * int(h[2:4], 16) + 0.114 * int(h[4:6], 16)
                return lum < 128
            except Exception:
                return False

        def fmt_value(cell):
            v = cell.value
            if v is None:
                return ""
            if isinstance(v, datetime):
                return v.strftime("%d/%m/%Y %H:%M") if (v.hour or v.minute) else v.strftime("%d/%m/%Y")
            if isinstance(v, date):
                return v.strftime("%d/%m/%Y")
            if isinstance(v, time):
                return v.strftime("%H:%M")
            if isinstance(v, bool):
                return "TRUE" if v else "FALSE"
            if isinstance(v, (int, float)):
                fmt = (cell.number_format or "").lower()
                try:
                    if "%" in fmt:
                        return f"{v * 100:g}%"
                    if "£" in fmt or "gbp" in fmt:
                        return f"£{v:,.2f}"
                    if "$" in fmt:
                        return f"${v:,.2f}"
                    if "," in fmt and isinstance(v, float):
                        return f"{v:,.2f}"
                    if "," in fmt:
                        return f"{v:,}"
                except Exception:
                    pass
                if isinstance(v, float) and v.is_integer():
                    return str(int(v))
                return str(v)
            return str(v)

        def emu_px(e):
            return round((e or 0) / 9525)

        def images_html(ws, col_px, row_px):
            """Embedded pictures, absolutely positioned over the sheet from their
            anchor cell. Best-effort — any image that can't be placed is skipped."""
            out = []
            for image in list(getattr(ws, "_images", []) or [])[:20]:
                try:
                    raw = None
                    d = getattr(image, "_data", None)
                    if callable(d):
                        raw = d()
                    elif isinstance(d, (bytes, bytearray)):
                        raw = bytes(d)
                    if raw is None and hasattr(image, "ref"):
                        ref = image.ref
                        if hasattr(ref, "read"):
                            ref.seek(0); raw = ref.read()
                        elif hasattr(ref, "save"):
                            buf = io.BytesIO(); ref.save(buf, format="PNG"); raw = buf.getvalue()
                    if not raw:
                        continue
                    ext = (getattr(image, "format", None) or "png").lower()
                    mime = "jpeg" if ext in ("jpg", "jpeg") else ext
                    b64 = base64.b64encode(raw).decode("ascii")
                    anc = image.anchor
                    frm = getattr(anc, "_from", None)
                    if frm is None:
                        continue
                    left = ROWHDR_W + sum(col_px[:frm.col]) + emu_px(getattr(frm, "colOff", 0))
                    top = HROW_PX + sum(row_px[:frm.row]) + emu_px(getattr(frm, "rowOff", 0))
                    to = getattr(anc, "to", None)
                    if to is not None:
                        w = max(8, ROWHDR_W + sum(col_px[:to.col]) + emu_px(getattr(to, "colOff", 0)) - left)
                        h = max(8, HROW_PX + sum(row_px[:to.row]) + emu_px(getattr(to, "rowOff", 0)) - top)
                    else:
                        ext_o = getattr(anc, "ext", None)
                        w = emu_px(getattr(ext_o, "cx", 0)) or emu_px(getattr(image, "width", 0)) or 96
                        h = emu_px(getattr(ext_o, "cy", 0)) or emu_px(getattr(image, "height", 0)) or 96
                    out.append(
                        f"<img src='data:image/{mime};base64,{b64}' "
                        f"style='position:absolute;left:{left}px;top:{top}px;width:{w}px;"
                        f"height:{h}px;z-index:0;object-fit:contain'>")
                except Exception:
                    continue
            return "".join(out)

        HDR = ("position:sticky;top:0;z-index:4;background:#f3f3f3;border:1px solid #c8c8c8;"
               "padding:3px 6px;font-weight:600;color:#555;text-align:center")
        RHDR = ("position:sticky;left:0;z-index:3;background:#f3f3f3;border:1px solid #c8c8c8;"
                "padding:3px 6px;color:#777;text-align:center;font-weight:500")
        CORNER = ("position:sticky;top:0;left:0;z-index:5;background:#e9e9e9;border:1px solid #c8c8c8")

        parts: List[str] = []
        for ws in wb.worksheets:
            max_row = min(ws.max_row or 1, MAX_ROWS)
            max_col = min(ws.max_column or 1, MAX_COLS)

            # Merged cells → colspan/rowspan on the anchor; skip the covered cells.
            anchors, covered = {}, set()
            for rng in ws.merged_cells.ranges:
                anchors[(rng.min_row, rng.min_col)] = (rng.max_row - rng.min_row + 1,
                                                       rng.max_col - rng.min_col + 1)
                for rr in range(rng.min_row, rng.max_row + 1):
                    for cc in range(rng.min_col, rng.max_col + 1):
                        if (rr, cc) != (rng.min_row, rng.min_col):
                            covered.add((rr, cc))

            col_px = []
            for c in range(1, max_col + 1):
                w = None
                try:
                    cd = ws.column_dimensions.get(get_column_letter(c))
                    if cd and cd.width:
                        w = int(cd.width * 7 + 6)
                except Exception:
                    w = None
                col_px.append(w or DEF_COL_PX)
            row_px = []
            for r in range(1, max_row + 1):
                h = None
                try:
                    rd = ws.row_dimensions.get(r)
                    if rd and rd.height:
                        h = int(rd.height * 4 / 3)
                except Exception:
                    h = None
                row_px.append(h or DEF_ROW_PX)

            colgroup = [f"<col style='width:{ROWHDR_W}px'>"] + [
                f"<col style='width:{w}px'>" for w in col_px]

            header = [f"<th style='{CORNER}'></th>"]
            for c in range(1, max_col + 1):
                header.append(f"<th style='{HDR}'>{get_column_letter(c)}</th>")
            body_rows = [f"<tr style='height:{HROW_PX}px'>{''.join(header)}</tr>"]

            for r in range(1, max_row + 1):
                tds = [f"<th style='{RHDR}'>{r}</th>"]
                for c in range(1, max_col + 1):
                    if (r, c) in covered:
                        continue
                    cell = ws.cell(row=r, column=c)
                    span = anchors.get((r, c))
                    attrs = ""
                    if span:
                        if span[0] > 1:
                            attrs += f" rowspan={span[0]}"
                        if span[1] > 1:
                            attrs += f" colspan={span[1]}"
                    # Base faint gridline; explicit borders override each side.
                    bt = br = bb = bl = "1px solid #e7e7e7"
                    bd = cell.border
                    if bd is not None:
                        bt = side_border(bd.top) or bt
                        br = side_border(bd.right) or br
                        bb = side_border(bd.bottom) or bb
                        bl = side_border(bd.left) or bl
                    st = [f"border-top:{bt}", f"border-right:{br}", f"border-bottom:{bb}",
                          f"border-left:{bl}", "padding:2px 6px", "white-space:nowrap",
                          "overflow:hidden", "text-overflow:ellipsis", "max-width:360px"]
                    bg = None
                    try:
                        if cell.fill is not None and cell.fill.patternType == "solid":
                            bg = resolve_color(cell.fill.fgColor)
                            if bg and bg.upper() != "#FFFFFF":
                                st.append(f"background:{bg}")
                            else:
                                bg = None
                    except Exception:
                        bg = None
                    f = cell.font
                    fc = resolve_color(getattr(f, "color", None)) if f is not None else None
                    if f is not None:
                        if f.bold:
                            st.append("font-weight:600")
                        if f.italic:
                            st.append("font-style:italic")
                    if fc:
                        st.append(f"color:{fc}")
                    elif bg and is_dark(bg):
                        st.append("color:#ffffff")  # keep text legible on dark fills
                    al = cell.alignment.horizontal if cell.alignment else None
                    if al in ("center", "right", "left"):
                        st.append(f"text-align:{al}")
                    elif isinstance(cell.value, (int, float)) and not isinstance(cell.value, bool):
                        st.append("text-align:right")
                    tds.append(f"<td{attrs} style='{';'.join(st)}'>{_html.escape(fmt_value(cell))}</td>")
                body_rows.append(f"<tr style='height:{row_px[r - 1]}px'>{''.join(tds)}</tr>")

            imgs = images_html(ws, col_px, row_px)
            parts.append(f"<div style='font-weight:600;margin:10px 0 4px;color:#333'>{_html.escape(ws.title)}</div>")
            parts.append(
                "<div style='position:relative;display:inline-block'>"
                "<table style='border-collapse:collapse;table-layout:fixed;font-size:12px;color:#1a1a1a'>"
                f"<colgroup>{''.join(colgroup)}</colgroup>{''.join(body_rows)}</table>"
                + imgs + "</div>"
            )
            if (ws.max_row or 0) > MAX_ROWS or (ws.max_column or 0) > MAX_COLS:
                parts.append("<div style='color:#999;font-size:11px;margin:4px 0 10px'>"
                             f"Showing first {max_row} rows × {max_col} columns.</div>")
        wb.close()
        return CaseHistoryService._grid_wrap("".join(parts) or "<p style='padding:12px'>Empty workbook.</p>")

    @staticmethod
    def _xls_to_html(data: bytes) -> str:
        """Render a legacy .xls workbook as an Excel-style HTML grid via xlrd (values
        only — the old binary format's styling isn't read here)."""
        import html as _html
        import xlrd

        book = xlrd.open_workbook(file_contents=data)
        MAX_ROWS, MAX_COLS = 500, 60
        from openpyxl.utils import get_column_letter

        HDR = ("position:sticky;top:0;z-index:2;background:#f3f3f3;border:1px solid #c8c8c8;"
               "padding:3px 6px;font-weight:600;color:#555;text-align:center")
        RHDR = ("position:sticky;left:0;z-index:1;background:#f3f3f3;border:1px solid #c8c8c8;"
                "padding:3px 6px;color:#777;text-align:center;font-weight:500")
        CORNER = ("position:sticky;top:0;left:0;z-index:3;background:#e9e9e9;border:1px solid #c8c8c8")

        parts: List[str] = []
        for sh in book.sheets():
            nrows, ncols = min(sh.nrows, MAX_ROWS), min(sh.ncols, MAX_COLS)
            if nrows == 0 or ncols == 0:
                continue
            header = [f"<th style='{CORNER}'></th>"] + [
                f"<th style='{HDR}'>{get_column_letter(c + 1)}</th>" for c in range(ncols)]
            body_rows = [f"<tr>{''.join(header)}</tr>"]
            for r in range(nrows):
                tds = [f"<th style='{RHDR}'>{r + 1}</th>"]
                for c in range(ncols):
                    v = sh.cell_value(r, c)
                    if isinstance(v, float) and v.is_integer():
                        v = int(v)
                    align = "text-align:right" if isinstance(v, (int, float)) else "text-align:left"
                    tds.append(
                        f"<td style='border:1px solid #d7d7d7;padding:2px 6px;white-space:nowrap;"
                        f"max-width:360px;overflow:hidden;text-overflow:ellipsis;{align}'>"
                        f"{_html.escape('' if v == '' else str(v))}</td>")
                body_rows.append(f"<tr>{''.join(tds)}</tr>")
            parts.append(f"<div style='font-weight:600;margin:10px 0 4px;color:#333'>{_html.escape(sh.name)}</div>")
            parts.append("<table style='border-collapse:collapse;font-size:12px;color:#1a1a1a'>"
                         f"{''.join(body_rows)}</table>")
        return CaseHistoryService._grid_wrap("".join(parts) or "<p style='padding:12px'>Empty workbook.</p>")

    @staticmethod
    def _grid_wrap(body: str) -> str:
        """Wrap spreadsheet HTML in a scrollable, Excel-like framed box. The
        ``scrollbar-hide`` class (defined in the app's global CSS) + scrollbar-width
        keep the scrollbar invisible while it still scrolls."""
        return (
            "<div class=\"scrollbar-hide\" style=\"font-family:Calibri,'Segoe UI',Arial,"
            "sans-serif;max-height:68vh;overflow:auto;scrollbar-width:none;"
            "-ms-overflow-style:none;border:1px solid #d0d0d0;border-radius:4px;"
            "background:#fff;padding:0 4px 8px\">" + body + "</div>"
        )

    @staticmethod
    def claim_emails(db: Session, claim_id: int) -> List[dict]:
        """Emails matching this claim's case reference, shaped as read-only History
        records. See ``emails_by_reference``."""
        claim = db.query(Claim).filter(Claim.id == claim_id).first()
        if not claim:
            return []
        return CaseHistoryService.emails_by_reference(
            db, build_case_reference(claim.id, db), scope_type="claim", scope_id=claim_id,
        )

    @staticmethod
    def _normalize_subject(subject: Optional[str]) -> str:
        """Strip Re:/Fwd: prefixes + collapse whitespace so replies share a key."""
        import re
        s = re.sub(r"^(?:\s*(re|fw|fwd|aw|wg)\s*:\s*)+", "", (subject or "").strip(), flags=re.IGNORECASE)
        return re.sub(r"\s+", " ", s).strip().lower()

    @staticmethod
    def _thread_key(conversation_id: Optional[str], subject: Optional[str]) -> str:
        """Group emails into a thread: Outlook conversation id when present, else the
        normalized subject (so a reply lands in the same thread as the original)."""
        cid = (conversation_id or "").strip()
        if cid:
            return f"conv:{cid}"
        norm = CaseHistoryService._normalize_subject(subject)
        return f"subj:{norm}" if norm else ""

    @staticmethod
    def emails_by_reference(db: Session, reference: str, *, scope_type: str, scope_id: int) -> List[dict]:
        """Emails (sent + received) whose subject/body mention ``reference``, pulled
        live from the configured Outlook mailbox via Microsoft Graph, shaped as
        read-only History records (SE for outgoing, IE for incoming). Claims and
        Fleet share one mailbox but pass their own reference. Best-effort."""
        if not reference:
            return []
        try:
            token = MicrosoftGraphTokenService.get_access_token("read")
            if not token:
                return []
            items = OutlookCaseActivityService.get_case_emails(
                claim_reference=reference, access_token=token
            )
        except Exception as exc:  # never let a mailbox hiccup break the History screen
            print(f"[CaseHistoryService] email fetch error: {exc}")
            return []

        # Our own addresses (read mailbox + the noreply we send from) — the
        # correspondent is the OTHER party, never one of these.
        import os
        ours = {"no-replynationwideassist@outlook.com"}
        mb = (MicrosoftGraphTokenService.mailbox_user("read") or "").strip().lower()
        if mb:
            ours.add(mb)
        ours.add((os.getenv("SENDGRID_SENDER") or "no-replynationwideassist@outlook.com").strip().lower())

        def _correspondent(item) -> Optional[str]:
            frm = (item.sender_email or "").strip()
            to_list = [a for a in ((item.meta or {}).get("to_recipients") or []) if a]
            if frm.lower() in ours:
                # Sent by us (e.g. from the noreply) → correspondent = the recipient.
                for a in to_list:
                    if a.strip().lower() not in ours:
                        return a
                return to_list[0] if to_list else (frm or None)
            # Received → the sender is the correspondent.
            return frm or (to_list[0] if to_list else None)

        out: List[dict] = []
        for i, it in enumerate(items):
            attachments = [
                {"name": a.file_name, "url": a.file_url, "size": a.file_size}
                for a in (it.attachments or [])
            ]
            # Everything fetched from the Outlook mailbox is treated as Incoming Email
            # (IE) — these are messages that landed in the mailbox. Outgoing (SE) records
            # are only created when the user replies/forwards from the app.
            out.append({
                "id": f"email:{i}",
                "claim_id": scope_id if scope_type == "claim" else None,
                "scope_type": scope_type,
                "scope_id": scope_id,
                "action_type": "incoming_email",
                "posted_at": it.received_at,
                "correspondent": _correspondent(it),
                "handler": it.sender_name or None,
                "subject": it.subject or None,
                "details": it.body_preview or None,
                "payload": {
                    "source": "email",
                    # Graph message id — lets the History detail pane reply/forward
                    # this email through the same Case Activity email endpoints.
                    "message_id": (it.meta or {}).get("message_id") or it.id or None,
                    # Outlook conversation id groups messages into a thread; the frontend
                    # groups by thread_key (conversation id, else the normalized subject).
                    "conversation_id": (it.meta or {}).get("conversation_id") or None,
                    "thread_key": CaseHistoryService._thread_key(
                        (it.meta or {}).get("conversation_id"), it.subject),
                    "from_name": it.sender_name,
                    "from_email": it.sender_email,
                    "to": (it.meta or {}).get("to_recipients") or [],
                    "body_html": it.body_html,
                    "body_text": it.body_text,
                    "attachments": attachments,
                },
                "created_by": None,
                "created_at": it.received_at,
            })
        return out

    # Reference prefixes staff put in the subject/body so the shared mailbox can be
    # filtered per fleet entity (claims use the case reference).
    _SCOPE_PREFIX = {"fleet_hire": "FH", "vm_cams": "CAMS", "vm_skyline": "SKY"}

    @staticmethod
    def scope_reference(db: Session, scope_type: str, scope_id: int) -> str:
        if scope_type == "claim":
            return build_case_reference(scope_id, db)
        prefix = CaseHistoryService._SCOPE_PREFIX.get(scope_type, scope_type.upper())
        return f"{prefix}-{scope_id}"

    @staticmethod
    def scope_emails(db: Session, scope_type: str, scope_id: int) -> List[dict]:
        if scope_type == "claim":
            return CaseHistoryService.claim_emails(db, scope_id)
        return CaseHistoryService.emails_by_reference(
            db, CaseHistoryService.scope_reference(db, scope_type, scope_id),
            scope_type=scope_type, scope_id=scope_id,
        )

    @staticmethod
    def filter_options(db: Session, claim_id: int) -> dict:
        return CaseHistoryService.filter_options_for_scope(db, "claim", claim_id)

    @staticmethod
    def filter_options_for_scope(db: Session, scope_type: str, scope_id: int) -> dict:
        rows = (
            db.query(CaseHistory)
            .filter(_scope_clause(scope_type, scope_id), CaseHistory.is_deleted.isnot(True))
            .all()
        )
        return {
            "correspondents": sorted({r.correspondent for r in rows if r.correspondent}),
            "handlers": sorted({r.handler for r in rows if r.handler}),
            "action_types": sorted({r.action_type.value for r in rows if r.action_type}),
        }
