"""Generated Fleet document assets.

These files live inside the Fleet module so Fleet stays separable from Claims.
For now the Raise Hire Documentation source files are static Office documents
from ``fleet/assets/Documents``; this service exposes them through authenticated
Fleet routes for download/email attachment.
"""
from dataclasses import dataclass
from datetime import date
from html import escape
from io import BytesIO
import mimetypes
from pathlib import Path
import re
import sys
from typing import Dict, List, Optional, Tuple
from zipfile import ZIP_DEFLATED, ZipFile

from fastapi import HTTPException
from sqlalchemy.orm import Session

from fleet.deps import S3Service
from fleet.models.tables import FleetHireDocument, FleetHireVehicle
from fleet.services.common import get_hire_or_404
from fleet.services import ocr as fleet_ocr


ASSET_DIR = Path(__file__).resolve().parents[1] / "assets" / "Documents"


@dataclass(frozen=True)
class GeneratedDocumentAsset:
    key: str
    filename: str
    source_filename: str

    @property
    def path(self) -> Path:
        return ASSET_DIR / self.source_filename

    @property
    def content_type(self) -> str:
        return mimetypes.guess_type(self.filename)[0] or "application/octet-stream"


DOCUMENT_GROUPS: Dict[str, List[GeneratedDocumentAsset]] = {
    "raise_hire_documentation": [
        GeneratedDocumentAsset(
            key="raise_hire_documentation_docx",
            filename="Raise Hire Documentation.docx",
            source_filename="Raise Hire Documentation .docx",
        ),
        GeneratedDocumentAsset(
            key="raise_hire_documentation_xls",
            filename="Raise Hire Documentation II.xls",
            source_filename="Raise Hire Documentation II.xls",
        ),
    ],
    "raise_authority_letter": [
        GeneratedDocumentAsset(
            key="raise_authority_letter_docx",
            filename="Raise Authority Letter.docx",
            source_filename="Raise Authority Letter.docx",
        ),
    ],
    "raise_vehicle_inspection_sheet": [
        GeneratedDocumentAsset(
            key="raise_vehicle_inspection_sheet_xlsx",
            filename="Vehicle Inspection Sheet.xlsx",
            source_filename="Vehicle Inspection Sheet.xlsx",
        ),
    ],
}


def _assets_for(document_key: str) -> List[GeneratedDocumentAsset]:
    assets = DOCUMENT_GROUPS.get(document_key)
    if not assets:
        raise HTTPException(status_code=404, detail="Generated document not found")
    missing = [asset.filename for asset in assets if not asset.path.exists()]
    if missing:
        raise HTTPException(status_code=404, detail=f"Document asset missing: {', '.join(missing)}")
    return assets


def list_document_files(
    db: Session,
    hire_id: int,
    tenant_id: Optional[int],
    document_key: str,
    vehicle_id: Optional[int] = None,
) -> List[dict]:
    hire, vehicle = _get_hire_and_vehicle(db, hire_id, tenant_id, vehicle_id)
    context = _document_context(hire, vehicle, db)
    return [
        {
            "key": asset.key,
            "filename": asset.filename,
            "content_type": asset.content_type,
            "size": len(_render_asset(asset, context)),
        }
        for asset in _assets_for(document_key)
    ]


def get_document_file(
    db: Session,
    hire_id: int,
    tenant_id: Optional[int],
    document_key: str,
    file_key: str,
    vehicle_id: Optional[int] = None,
) -> Tuple[bytes, str, str]:
    hire, vehicle = _get_hire_and_vehicle(db, hire_id, tenant_id, vehicle_id)
    asset = next((item for item in _assets_for(document_key) if item.key == file_key), None)
    if not asset:
        raise HTTPException(status_code=404, detail="Generated document file not found")
    return _render_asset(asset, _document_context(hire, vehicle, db)), asset.content_type, asset.filename


def get_document_bundle(
    db: Session,
    hire_id: int,
    tenant_id: Optional[int],
    document_key: str,
    vehicle_id: Optional[int] = None,
) -> Tuple[bytes, str, str]:
    hire, vehicle = _get_hire_and_vehicle(db, hire_id, tenant_id, vehicle_id)
    context = _document_context(hire, vehicle, db)
    assets = _assets_for(document_key)
    if len(assets) == 1:
        asset = assets[0]
        return _render_asset(asset, context), asset.content_type, asset.filename

    output = BytesIO()
    with ZipFile(output, "w", compression=ZIP_DEFLATED) as zf:
        for asset in assets:
            zf.writestr(asset.filename, _render_asset(asset, context))
    return output.getvalue(), "application/zip", f"{document_key.replace('_', ' ').title()}.zip"


def get_document_print_view(
    db: Session,
    hire_id: int,
    tenant_id: Optional[int],
    document_key: str,
    vehicle_id: Optional[int] = None,
) -> str:
    hire, vehicle = _get_hire_and_vehicle(db, hire_id, tenant_id, vehicle_id)
    context = _document_context(hire, vehicle, db)
    sections = []
    for asset in _assets_for(document_key):
        data = _render_asset(asset, context)
        if asset.filename.lower().endswith(".docx"):
            body = _docx_to_html(data)
        elif asset.filename.lower().endswith(".xls"):
            body = _xls_to_html(data)
        elif asset.filename.lower().endswith(".xlsx"):
            body = _xlsx_to_html(data)
        else:
            body = f"<p>This document is ready: {escape(asset.filename)}</p>"
        sections.append(f"<section><h2>{escape(asset.filename)}</h2>{body}</section>")

    return f"""<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <title>Printable Documents</title>
    <style>
      *{{box-sizing:border-box}}
      body{{font-family:Arial,sans-serif;margin:24px;color:#111827;background:#fff}}
      header{{display:flex;align-items:center;justify-content:space-between;border-bottom:1px solid #e5e7eb;padding-bottom:14px;margin-bottom:20px}}
      h1{{font-size:22px;margin:0}}
      h2{{font-size:16px;margin:0 0 14px}}
      button{{font:inherit;border:1px solid #111827;border-radius:4px;background:#111827;color:#fff;padding:10px 14px;cursor:pointer}}
      section{{break-after:page;margin-bottom:28px}}
      p{{margin:0 0 8px;line-height:1.45;white-space:pre-wrap}}
      table{{border-collapse:collapse;width:100%;margin:8px 0 14px;table-layout:auto}}
      td,th{{border:1px solid #d1d5db;padding:6px 8px;vertical-align:top;font-size:12px;white-space:pre-wrap}}
      th{{background:#f3f4f6;text-align:left}}
      .empty{{color:#9ca3af}}
      @media print{{
        body{{margin:0}}
        header{{display:none}}
        section{{margin:0 0 18px}}
      }}
    </style>
  </head>
  <body>
    <header><h1>Printable Documents</h1><button onclick="window.print()">Print</button></header>
    {''.join(sections)}
  </body>
</html>"""


def _get_hire_and_vehicle(db: Session, hire_id: int, tenant_id: Optional[int], vehicle_id: Optional[int]):
    hire = get_hire_or_404(db, hire_id, tenant_id)
    query = db.query(FleetHireVehicle).filter(FleetHireVehicle.hire_id == hire_id)
    if vehicle_id:
        query = query.filter(FleetHireVehicle.id == vehicle_id)
    vehicle = query.order_by(FleetHireVehicle.position, FleetHireVehicle.id).first()
    if vehicle_id and not vehicle:
        raise HTTPException(status_code=404, detail="Hire vehicle not found")
    return hire, vehicle


def _docx_to_html(data: bytes) -> str:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(BytesIO(data))

    def paragraph_html(paragraph: Paragraph) -> str:
        text = paragraph.text.strip()
        if not text:
            return ""
        style = (paragraph.style.name if paragraph.style else "").lower()
        tag = "h3" if "heading" in style or text.isupper() and len(text) < 80 else "p"
        return f"<{tag}>{escape(text)}</{tag}>"

    def table_html(table: Table) -> str:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{escape(cell.text.strip()) or '&nbsp;'}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        return f"<table>{''.join(rows)}</table>"

    parts = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            html = paragraph_html(Paragraph(child, document))
            if html:
                parts.append(html)
        elif isinstance(child, CT_Tbl):
            parts.append(table_html(Table(child, document)))
    return "".join(parts) or "<p class=\"empty\">No printable content found.</p>"


def _xlsx_to_html(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(data), data_only=True)
    sheet = workbook.active
    used_rows = []
    for row in sheet.iter_rows():
        values = [cell.value for cell in row]
        if any(value not in (None, "") for value in values):
            used_rows.append(values)
    if not used_rows:
        return "<p class=\"empty\">No printable content found.</p>"

    last_col = max(
        (idx + 1 for values in used_rows for idx, value in enumerate(values) if value not in (None, "")),
        default=0,
    )
    rows = []
    for values in used_rows:
        cells = "".join(
            f"<td>{escape(str(value)) if value not in (None, '') else '&nbsp;'}</td>"
            for value in values[:last_col]
        )
        rows.append(f"<tr>{cells}</tr>")
    return f"<table>{''.join(rows)}</table>"


def _xls_to_html(data: bytes) -> str:
    import xlrd

    workbook = xlrd.open_workbook(file_contents=data)
    sheet = workbook.sheet_by_index(0)
    rows = []
    for row_index in range(sheet.nrows):
        values = [sheet.cell_value(row_index, col_index) for col_index in range(sheet.ncols)]
        if any(value not in (None, "") for value in values):
            rows.append(values)
    if not rows:
        return "<p class=\"empty\">No printable content found.</p>"

    last_col = max(
        (idx + 1 for values in rows for idx, value in enumerate(values) if value not in (None, "")),
        default=0,
    )
    html_rows = []
    for values in rows:
        cells = "".join(
            f"<td>{escape(str(value)) if value not in (None, '') else '&nbsp;'}</td>"
            for value in values[:last_col]
        )
        html_rows.append(f"<tr>{cells}</tr>")
    return f"<table>{''.join(html_rows)}</table>"


def _format_date(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%d/%m/%Y")
    except AttributeError:
        return str(value)


def _display_date_from_dmy(value: str) -> str:
    match = re.match(r"^(\d{1,2})[-/](\d{1,2})[-/](\d{2,4})$", str(value or "").strip())
    if not match:
        return str(value or "").strip()
    day, month, year = match.groups()
    return f"{day.zfill(2)}/{month.zfill(2)}/{year if len(year) == 4 else '20' + year}"


def _date_time(value, time_value) -> str:
    parts = [_format_date(value), str(time_value or "").strip()]
    return " ".join(part for part in parts if part)


def _money_number(value) -> Optional[float]:
    cleaned = re.sub(r"[^0-9.\-]", "", str(value or ""))
    if not cleaned:
        return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def _format_money(value) -> str:
    number = _money_number(value)
    if number is None:
        return ""
    return f"£{number:.2f}"


def _format_money_plain(value) -> str:
    number = _money_number(value)
    if number is None:
        return ""
    return f"{number:.2f}"


def _address_line(hire) -> str:
    return ", ".join(part.strip() for part in [hire.driver_address, hire.driver_postcode] if part and str(part).strip())


def _licence_expiry_from_saved_doc(db: Optional[Session], hire_id: int) -> str:
    if db is None:
        return ""
    doc = (
        db.query(FleetHireDocument)
        .filter(
            FleetHireDocument.hire_id == hire_id,
            FleetHireDocument.doc_type == "dlFront",
        )
        .order_by(FleetHireDocument.id.desc())
        .first()
    )
    if not doc or not doc.s3_key:
        return ""
    try:
        data = S3Service().read_file_bytes(doc.s3_key)
        text = fleet_ocr.file_to_text(data, doc.filename or "")
        parsed = fleet_ocr.parse_driving_licence(text)
        return _display_date_from_dmy(parsed.get("licenceEnd", ""))
    except Exception as exc:  # pylint: disable=broad-exception-caught
        print(f"Fleet generated docs: licence expiry OCR fallback failed: {exc}")
        return ""


def _document_context(hire, vehicle: Optional[FleetHireVehicle], db: Optional[Session] = None) -> dict:
    weekly = (
        getattr(vehicle, "vehicle_cost_per_week", None)
        or getattr(hire, "weekly_hire_payment", None)
        or getattr(hire, "vehicle_cost_per_week", None)
    )
    deposit = (
        getattr(vehicle, "deposit", None)
        or getattr(hire, "security_deposit", None)
        or getattr(hire, "deposit", None)
    )
    total_due_number = (_money_number(weekly) or 0) + (_money_number(deposit) or 0)
    start_date = getattr(vehicle, "hire_start_date", None) or getattr(hire, "payment_hire_start_date", None) or getattr(hire, "hire_start_date", None)
    end_date = (
        getattr(vehicle, "hire_end_date", None)
        or getattr(vehicle, "checkout_date", None)
        or getattr(hire, "payment_hire_end_date", None)
        or getattr(hire, "hire_end_date", None)
    )
    start_time = getattr(vehicle, "hire_start_time", None)
    end_time = getattr(vehicle, "checkout_time", None)
    make = (getattr(vehicle, "make", None) or getattr(hire, "make", None) or "").strip()
    model = (getattr(vehicle, "model", None) or getattr(hire, "model", None) or "").strip()
    licence_expiry = _format_date(hire.driving_licence_end) or _licence_expiry_from_saved_doc(db, hire.id)

    return {
        "document_date": date.today().strftime("%d/%m/%Y"),
        "fleet_reference": (hire.fleet_reference or "").strip(),
        "hirer_name": (hire.driver_name or "").strip(),
        "hirer_address": (hire.driver_address or "").strip(),
        "hirer_postcode": (hire.driver_postcode or "").strip(),
        "hirer_full": " of ".join(part for part in [(hire.driver_name or "").strip(), _address_line(hire)] if part),
        "date_of_birth": _format_date(hire.date_of_birth),
        "driving_licence_number": (hire.driving_licence_number or "").strip(),
        "licence_expiry": licence_expiry,
        "make": make,
        "model": model,
        "vehicle_description": " ".join(part for part in [make, model] if part),
        "registration": (getattr(vehicle, "registration_number", None) or getattr(hire, "registration_number", None) or "").strip(),
        "hire_start": _date_time(start_date, start_time),
        "hire_start_date": _format_date(start_date),
        "hire_start_time": str(start_time or "").strip(),
        "hire_end": _date_time(end_date, end_time),
        "hire_end_date": _format_date(end_date),
        "hire_end_time": str(end_time or "").strip(),
        "mileage_start": str(getattr(vehicle, "mileage_start", None) or "").strip(),
        "mileage_end": str(getattr(vehicle, "mileage_end", None) or "").strip(),
        "weekly": _format_money(weekly),
        "weekly_plain": _format_money_plain(weekly),
        "weekly_number": _money_number(weekly),
        "deposit": _format_money(deposit),
        "deposit_plain": _format_money_plain(deposit),
        "deposit_number": _money_number(deposit),
        "total_due": f"£{total_due_number:.2f}" if total_due_number else "",
        "total_due_number": total_due_number if total_due_number else None,
    }


def _set_paragraph_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_sample_text(text: str, ctx: dict) -> str:
    address = ", ".join(part for part in [ctx["hirer_address"], ctx["hirer_postcode"]] if part)
    replacements = {
        "Mr SabER Mehrabi": ctx["hirer_name"],
        "Mona Lisa": ctx["hirer_name"],
        "54 Highfields court t , Leasowes Drive , WV44PZ": address,
        "54 Highfields court t , Leasowes Drive": ctx["hirer_address"],
        "54 Highfields court t \nLeasowes Drive": ctx["hirer_address"],
        "9 Anderson Drive Aberdeen Pea, AB15 4ST": address,
        "WV44PZ": ctx["hirer_postcode"],
        "MEHRA901276S99ZM 49": ctx["driving_licence_number"],
        "LISA753116SM9IJ": ctx["driving_licence_number"],
        "JK20XYZ": ctx["registration"],
        "BMW 1 SERIES SE": ctx["vehicle_description"],
        "13/07/26": ctx["hire_start_date"],
        "13/07/2026": ctx["hire_start"],
        "03/02/2026": ctx["hire_end"],
        "26/01/2066": ctx["licence_expiry"],
        "£300.00": ctx["deposit"],
        "£230.00": ctx["weekly"],
    }
    updated = text
    for old, new in replacements.items():
        if new:
            updated = updated.replace(old, new)
    if ctx["registration"]:
        updated = re.sub(r"vehicle registration\s*,", f"vehicle registration {ctx['registration']},", updated, flags=re.I)
        updated = re.sub(
            r"vehicle registration\s+from Skyline",
            f"vehicle registration {ctx['registration']} from Skyline",
            updated,
            flags=re.I,
        )
    return updated


def _render_docx(asset: GeneratedDocumentAsset, ctx: dict) -> bytes:
    from docx import Document

    doc = Document(asset.path)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Date:") and "Company:" in text:
            _set_paragraph_text(
                paragraph,
                re.sub(r"Date:\s*\d{1,2}/\d{1,2}/\d{2,4}", f"Date:{ctx['document_date']}", paragraph.text),
            )
        elif text.startswith("Hirer/Lessee"):
            _set_paragraph_text(paragraph, f"Hirer/Lessee (you/your) - \t{ctx['hirer_full']}")
        elif text.startswith("Full Name:") and "of" in text:
            _set_paragraph_text(paragraph, f"Full Name: \t\t{ctx['hirer_full']}")
        elif text.startswith("Full Name:"):
            _set_paragraph_text(paragraph, f"Full Name: {ctx['hirer_name']}")
        elif text.startswith("Driver Name:"):
            _set_paragraph_text(paragraph, f"Driver Name:\t\t\t{ctx['hirer_name']}")
        elif text.startswith("Address:") and "Church Lane" not in text:
            _set_paragraph_text(paragraph, f"Address:\t\t\t{', '.join(part for part in [ctx['hirer_address'], ctx['hirer_postcode']] if part)}")
        elif text.startswith("Driving License Number:"):
            _set_paragraph_text(paragraph, f"Driving License Number:\t{ctx['driving_licence_number']}")
        elif text.startswith("Vehicle Registration:"):
            _set_paragraph_text(paragraph, f"Vehicle Registration:\t{ctx['registration']}")
        elif text.startswith("Make & Model:"):
            _set_paragraph_text(paragraph, f"Make & Model:\t\t{ctx['vehicle_description']}")
        elif text.startswith("Hire Start Date:"):
            _set_paragraph_text(paragraph, f"Hire Start Date: {ctx['hire_start']} \tHire End Date: {ctx['hire_end']}")
        elif text.startswith("Vehicle Description (Make/Model)"):
            _set_paragraph_text(paragraph, f"Vehicle Description (Make/Model) -\t{ctx['vehicle_description']}")
        elif text.startswith("Vehicle Registration Number"):
            _set_paragraph_text(paragraph, f"Vehicle Registration Number -\t{ctx['registration']}")
        elif text.startswith("Agreement Start Date"):
            _set_paragraph_text(paragraph, f"Agreement Start Date -\t{ctx['hire_start']}")
        elif text.startswith("Expiry Date:"):
            _set_paragraph_text(paragraph, f"Expiry Date: \t\t{ctx['licence_expiry']}")
        elif text.startswith("Security Deposit"):
            _set_paragraph_text(paragraph, f"Security Deposit -\t\t\t\t{ctx['deposit']}")
        elif text.startswith("Payment Profile"):
            _set_paragraph_text(paragraph, f"Payment Profile -\tWeekly payments of {ctx['weekly']} Inc VAT/IPT")
        elif text.startswith("Hirer Name:"):
            _set_paragraph_text(paragraph, f"Hirer Name: {ctx['hirer_name']} \tDate: {ctx['hire_start']}")
        else:
            updated = _replace_sample_text(paragraph.text, ctx)
            if updated != paragraph.text:
                _set_paragraph_text(paragraph, updated)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def render_raise_authority_letter_docx(hire, vehicle: Optional[FleetHireVehicle], db: Optional[Session] = None) -> bytes:
    """Render the manager-provided Raise Authority Letter template as a DOCX."""
    asset = DOCUMENT_GROUPS["raise_authority_letter"][0]
    return _render_asset(asset, _document_context(hire, vehicle, db))


def _existing_xls_style_idx(ws, row: int, col: int):
    row_obj = getattr(ws, "_Worksheet__rows", {}).get(row)
    cell = getattr(row_obj, "_Row__cells", {}).get(col) if row_obj else None
    return getattr(cell, "xf_idx", None)


def _write_xls(ws, row: int, col: int, value):
    if value is not None:
        xf_idx = _existing_xls_style_idx(ws, row, col)
        ws.write(row, col, value)
        if xf_idx is not None:
            row_obj = getattr(ws, "_Worksheet__rows", {}).get(row)
            cell = getattr(row_obj, "_Row__cells", {}).get(col) if row_obj else None
            if cell is not None:
                cell.xf_idx = xf_idx


def _ensure_xls_template_packages():
    try:
        import xlrd  # noqa: F401
        import xlutils.copy  # noqa: F401
        return
    except ModuleNotFoundError:
        pass

    repo_root = Path(__file__).resolve().parents[3]
    local_site_packages = (
        repo_root
        / ".venv"
        / "lib"
        / f"python{sys.version_info.major}.{sys.version_info.minor}"
        / "site-packages"
    )
    if local_site_packages.exists():
        sys.path.insert(0, str(local_site_packages))

    import xlrd  # noqa: F401
    import xlutils.copy  # noqa: F401


def _render_hire_documentation_xls(asset: GeneratedDocumentAsset, ctx: dict) -> bytes:
    _ensure_xls_template_packages()
    import xlrd
    from xlutils.copy import copy

    book = xlrd.open_workbook(str(asset.path), formatting_info=True)
    writable = copy(book)
    ws = writable.get_sheet(0)

    _write_xls(ws, 1, 5, ctx["fleet_reference"])
    _write_xls(ws, 10, 4, ctx["hirer_name"])
    _write_xls(ws, 12, 4, ctx["hirer_address"])
    _write_xls(ws, 18, 4, ctx["hirer_postcode"])
    _write_xls(ws, 20, 4, ctx["date_of_birth"])
    _write_xls(ws, 22, 4, ctx["driving_licence_number"])
    _write_xls(ws, 25, 5, ctx["licence_expiry"])

    _write_xls(ws, 18, 9, ctx["make"])
    _write_xls(ws, 18, 13, ctx["model"])
    _write_xls(ws, 20, 9, ctx["registration"])
    _write_xls(ws, 22, 9, ctx["hire_start_date"])
    _write_xls(ws, 22, 13, ctx["hire_start_time"])
    _write_xls(ws, 25, 9, ctx["hire_end_date"])
    _write_xls(ws, 25, 13, ctx["hire_end_time"])

    # Additional driver section stays blank until the business confirms that flow.
    for row, col in [(34, 4), (40, 4), (51, 4), (53, 4), (55, 4), (58, 5), (61, 4)]:
        _write_xls(ws, row, col, "")
    _write_xls(ws, 34, 11, ctx["weekly_number"])
    _write_xls(ws, 37, 11, ctx["deposit_number"])
    _write_xls(ws, 40, 11, ctx["total_due_number"])

    for row, col in [(77, 12), (83, 5), (100, 5), (100, 12), (110, 10)]:
        _write_xls(ws, row, col, ctx["hire_start_date"])
    _write_xls(ws, 110, 6, ctx["hirer_name"])

    output = BytesIO()
    writable.save(output)
    return output.getvalue()


def _render_vehicle_inspection_xlsx(asset: GeneratedDocumentAsset, ctx: dict) -> bytes:
    from openpyxl import load_workbook

    workbook = load_workbook(asset.path)
    sheet = workbook.active
    sheet["N3"] = ctx["fleet_reference"]
    sheet["D5"] = ctx["hirer_name"]
    sheet["D7"] = ctx["hirer_address"]
    sheet["D13"] = ctx["hirer_postcode"]
    sheet["D19"] = ctx["vehicle_description"]
    sheet["K19"] = ctx["registration"]
    sheet["C23"] = ctx["hire_start_date"]
    sheet["C25"] = ctx["hire_start_time"]
    sheet["J23"] = ctx["hire_end_date"]
    sheet["J25"] = ctx["hire_end_time"]
    # "Mileage:" label sits at B27; start mileage mirrors the start column (C),
    # end mileage the end column (J), like the date/time rows above.
    sheet["C27"] = ctx["mileage_start"]
    sheet["J27"] = ctx["mileage_end"]
    sheet["D80"] = ctx["hirer_name"]

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _render_asset(asset: GeneratedDocumentAsset, context: dict) -> bytes:
    if asset.source_filename.lower().endswith(".docx"):
        return _render_docx(asset, context)
    if asset.source_filename == "Raise Hire Documentation II.xls":
        return _render_hire_documentation_xls(asset, context)
    if asset.source_filename == "Vehicle Inspection Sheet.xlsx":
        return _render_vehicle_inspection_xlsx(asset, context)
    return asset.path.read_bytes()
