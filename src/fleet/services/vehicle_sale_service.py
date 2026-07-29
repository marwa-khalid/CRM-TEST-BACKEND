"""Vehicle sale documentation: Release of Liability and Sale Receipt.

Reproduces the employer-provided "Raise Release of Liability and Receipt" Word
template, filled from the purchaser and sale details held on the vehicle record.
The stored template is a legacy binary .doc; the endpoint returns Word-compatible
HTML that mirrors its plain Times New Roman layout.
"""
import base64
from datetime import date, datetime
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Optional

from fleet.models.tables import FleetVehicleRecord

EM_DASH = "—"
SELLER_FULL = "Central Accident Management Services Ltd &amp; Nationwide Assist Ltd"
SELLER_PRIMARY = "Central Accident Management Services Ltd"
# Plain-text seller (no HTML entity) for the Word document.
SELLER_FULL_DOC = "Central Accident Management Services Ltd & Nationwide Assist Ltd"
DOTS_DOC = "." * 78
LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "logo.png"
# The CAMS letterhead extracted from the original Word template — re-embedded in
# the generated .docx so the download keeps the same logo as the source file.
DOCX_LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "Documents" / "cams_logo.jpg"


def _logo_data_uri() -> str:
    try:
        return f"data:image/jpeg;base64,{base64.b64encode(DOCX_LOGO_PATH.read_bytes()).decode('ascii')}"
    except OSError:
        return ""


def _fmt_date(value: Optional[date]) -> str:
    return value.strftime("%d/%m/%Y") if value else EM_DASH


def _money(value: Optional[str]) -> str:
    raw = (value or "").strip().replace(",", "").replace("£", "")
    if not raw:
        return EM_DASH
    try:
        return f"£{float(raw):,.2f}"
    except ValueError:
        return value or EM_DASH


def _num(value: Optional[str]) -> Optional[float]:
    try:
        return float((value or "").strip().replace(",", "").replace("£", ""))
    except ValueError:
        return None


def _vat(exc: Optional[str], inc: Optional[str]) -> str:
    e, i = _num(exc), _num(inc)
    if e is not None and i is not None:
        return f"£{i - e:,.2f}"
    return EM_DASH


def _vehicle_line(record: FleetVehicleRecord) -> str:
    parts = [(record.make or "").strip(), (record.model or "").strip(), (record.variant or "").strip()]
    return " ".join(p for p in parts if p) or EM_DASH


def _salvage_vehicle_line(record: FleetVehicleRecord, reg: str) -> str:
    make = (record.make or "").strip() or EM_DASH
    model = " ".join(
        p for p in ((record.model or "").strip(), (record.variant or "").strip()) if p
    ) or EM_DASH
    return f"{make} | {model} | {reg}"


def build_sale_documents_html(record: FleetVehicleRecord, show_print_button: bool = True) -> str:
    """Release of Liability (page 1) + Receipt of Sale (page 2).

    Mirrors the employer's "Raise Release of Liability and Receipt" template:
    compact Times New Roman paragraphs, logo header, and the same receipt table
    structure.
    """
    now = datetime.now()
    today = now.strftime("%d/%m/%Y")
    time_now = now.strftime("%-I:%M %p").lower()

    reg = (record.registration_number or "").strip() or EM_DASH
    make = (record.make or "").strip() or EM_DASH
    vehicle = _vehicle_line(record)
    salvage_vehicle = _salvage_vehicle_line(record, reg)
    # The sale date drives the letter body; fall back to today if not yet set.
    sold_on = _fmt_date(record.vehicle_sold_on) if record.vehicle_sold_on else today
    sign_stamp = f"{sold_on} {now.strftime('%H:%M:%S')}"
    purchaser_name = (record.purchaser_name or "").strip() or EM_DASH

    money_exc = _money(record.sold_for_exc_vat)
    money_inc = _money(record.sold_for_inc_vat)
    money_vat = _vat(record.sold_for_exc_vat, record.sold_for_inc_vat)
    product_name = make if make != EM_DASH else vehicle
    logo_src = _logo_data_uri()
    logo_html = f'<div class="doc-header"><img src="{logo_src}" alt="Nationwide Assist" /></div>' if logo_src else ""

    address_lines = [p.strip() for p in (record.purchaser_address or "").split(",") if p.strip()] or ["Purchaser address"]
    purchaser_block = "<br/>".join(
        escape(line) for line in [
            (record.purchaser_name or "").strip() or "Purchaser name",
            *address_lines,
            (record.purchaser_postcode or "").strip(),
        ] if line
    )

    def _sign_pair() -> str:
        return f"""
      <p class="sig"><strong>NAME: (BLOCK CAPITALS) {'.' * 82} &nbsp;&nbsp; DATE: {escape(sign_stamp)}</strong></p>
      <p class="sig-spacer"><br/></p>
      <p class="sig"><strong>SIGNED:</strong></p>
      <p class="sig-spacer"><br/></p>
      <p class="sig"><strong>Witnessed on behalf of {SELLER_FULL} by (Seller)</strong></p>
      <p class="sig-spacer"><br/></p>
      <p class="sig"><strong>NAME: (BLOCK CAPITALS) {'.' * 82} &nbsp;&nbsp; DATE: {escape(sign_stamp)}</strong></p>
      <p class="sig-spacer"><br/></p>
      <p class="sig"><strong>SIGNED:</strong></p>"""

    print_button = '<button class="screen-print" onclick="window.print()">Print</button>' if show_print_button else ""

    return f"""<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <title>Release of Liability and Receipt</title>
    <style>
      @page{{margin:72pt 72pt 72pt 72pt}}
      body{{font-family:"Times New Roman",Times,serif;margin:72pt;color:#000;background:#fff}}
      p{{margin:0 0 8px 0;line-height:16px;font-size:10px}}
      .top-space{{font-size:11px;line-height:16px;min-height:12px}}
      .doc-header{{text-align:center;margin:0 0 42px 0}}
      .doc-header img{{width:150px;height:auto;display:inline-block}}
      .blank{{min-height:11px}}
      .letter-body{{text-align:justify}}
      .sig{{font-size:9px;line-height:16px}}
      .sig-spacer{{font-size:9px;line-height:16px;min-height:10px}}
      .receipt-title{{font-size:14px;line-height:16px;font-weight:bold}}
      .receipt{{margin-top:132px}}
      table{{border-collapse:collapse;margin:16px 0 0 0;width:100%}}
      td{{font-size:9px;line-height:12px;vertical-align:top;padding:0 5px;border-color:#bfbfbf}}
      .head-cell{{border:1px solid #bfbfbf;font-size:11px;line-height:12px;font-weight:bold}}
      .inv{{width:100px;border:1px solid #bfbfbf}}
      .products{{border:1px solid #bfbfbf}}
      .blank-cell{{width:100px;border:0}}
      .label{{width:275px;border-top:1px solid #bfbfbf;border-right:1px solid #bfbfbf;text-align:right;font-weight:bold}}
      .amount{{width:115px;border:1px solid #bfbfbf}}
      .spacer-row td{{height:16px;border-top:0;border-bottom:0}}
      .screen-print{{position:fixed;right:24px;top:24px;font:14px Arial,sans-serif;border:1px solid #111;background:#111;color:#fff;padding:8px 12px;cursor:pointer}}
      @media print{{body{{margin:72pt}}.screen-print{{display:none}}}}
    </style>
  </head>
  <body>
    {print_button}
    <section class="release">
      {logo_html}
      <p>{purchaser_block}</p>
      <p class="blank"><br/></p>
      <p class="blank"><br/></p>
      <p>Our REF: {escape(reg)}</p>
      <p>Salvage Vehicle Purchased: {escape(salvage_vehicle)}</p>
      <p class="blank"><br/></p>
      <p>Date: {escape(sold_on)}</p>
      <p>Time: {escape(time_now)}</p>
      <p class="top-space"><br/></p>
      <p>Dear Sirs,</p>
      <p class="letter-body">
        I, {escape(purchaser_name)} sign to confirm that I have purchased the vehicle {escape(reg)}
        on the {escape(sold_on)} from the seller {SELLER_PRIMARY} and therefore release all
        liability for this vehicle from {SELLER_FULL}, as a result I {escape(purchaser_name)} now
        accept full liability for the vehicle mentioned above beginning the {escape(sold_on)} and
        confirm by signing below that I am the new owner of the vehicle {escape(reg)}.
      </p>
      <p>
        I also confirm that I will remove the vehicle purchased by me {escape(purchaser_name)} within 5 days of purchase and
        payment having been received by {SELLER_PRIMARY}/Nationwide Assist Ltd.
      </p>
      <p class="blank"><br/></p>
      {_sign_pair()}
    </section>

    <section class="receipt">
      {logo_html}
      <p class="receipt-title">RECEIPT OF SALE:</p>
      <p class="top-space"><br/></p>
      <p class="top-space"><br/></p>
      <p>Payee: {escape(purchaser_name)}</p>
      <p>Date: {escape(sold_on)}</p>
      <p class="top-space"><br/></p>
      <p class="top-space"><br/></p>
      <table cellspacing="0" cellpadding="0">
        <tr><td colspan="3" class="head-cell">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; INV NO. &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; Products</td></tr>
        <tr>
          <td class="inv">{escape(product_name)}</td>
          <td colspan="2" class="products">Sale of vehicle registration {escape(reg)}.</td>
        </tr>
        <tr>
          <td class="blank-cell"></td>
          <td class="label">Sub Total:</td>
          <td class="amount">{escape(money_exc)}</td>
        </tr>
        <tr class="spacer-row"><td class="blank-cell"></td><td class="label"></td><td class="amount"></td></tr>
        <tr>
          <td class="blank-cell"></td>
          <td class="label">VAT @ 20%</td>
          <td class="amount">{escape(money_vat)}</td>
        </tr>
        <tr class="spacer-row"><td class="blank-cell"></td><td class="label"></td><td class="amount"></td></tr>
        <tr>
          <td class="blank-cell"></td>
          <td class="label">Total:</td>
          <td class="amount">{escape(money_inc)}</td>
        </tr>
      </table>
      {_sign_pair()}
    </section>
  </body>
</html>"""


def _skyline_ref(record: FleetVehicleRecord) -> str:
    """Our REF on the sale documents is the Skyline hire reference (SK-HR-####),
    not the vehicle registration."""
    hire_id = getattr(record, "hire_id", None)
    return f"SK-HR-{hire_id:04d}" if hire_id else EM_DASH


def build_sale_documents_docx(record: FleetVehicleRecord) -> bytes:
    """Release of Liability + Receipt of Sale as a real Word .docx.

    Reproduces the employer's "Raise Release of Liability and Receipt" template
    — the CAMS letterhead, compact Times New Roman layout, and receipt table —
    filled from the record. Returned as an editable .docx (not HTML-as-.doc) so it
    opens with the same formatting and logo as the source file.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    now = datetime.now()
    today = now.strftime("%d/%m/%Y")
    time_now = now.strftime("%-I:%M %p").lower()

    reg = (record.registration_number or "").strip() or EM_DASH
    make = (record.make or "").strip() or EM_DASH
    vehicle = _vehicle_line(record)
    salvage_vehicle = _salvage_vehicle_line(record, reg)
    sold_on = _fmt_date(record.vehicle_sold_on) if record.vehicle_sold_on else today
    sign_stamp = f"{sold_on} {now.strftime('%H:%M:%S')}"
    purchaser_name = (record.purchaser_name or "").strip() or "Purchaser name"
    money_exc = _money(record.sold_for_exc_vat)
    money_inc = _money(record.sold_for_inc_vat)
    money_vat = _vat(record.sold_for_exc_vat, record.sold_for_inc_vat)
    product_name = make if make != EM_DASH else vehicle

    address_lines = [p.strip() for p in (record.purchaser_address or "").split(",") if p.strip()]
    if (record.purchaser_postcode or "").strip():
        address_lines.append(record.purchaser_postcode.strip())
    if not address_lines:
        address_lines = ["Purchaser address"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    def para(text="", *, bold=False, size=10, align=None, space_after=6):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"
        return p

    def logo():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        if DOCX_LOGO_PATH.exists():
            p.add_run().add_picture(str(DOCX_LOGO_PATH), width=Inches(1.2))

    def sign_pair():
        para(f"NAME: (BLOCK CAPITALS)    DATE: {sign_stamp}", bold=True, size=8.5, space_after=12)
        para("SIGNED:", bold=True, size=9, space_after=14)
        para(f"Witnessed on behalf of {SELLER_FULL_DOC} by (Seller)", bold=True, size=9, space_after=10)
        para(f"NAME: (BLOCK CAPITALS)    DATE: {sign_stamp}", bold=True, size=8.5, space_after=12)
        para("SIGNED:", bold=True, size=9, space_after=6)

    # ---- Page 1: Release of Liability ----
    logo()
    para(purchaser_name, space_after=0)
    for line in address_lines:
        para(line, space_after=0)
    para("", space_after=6)
    para(f"Our REF: {_skyline_ref(record)}", space_after=0)
    para(f"Salvage Vehicle Purchased: {salvage_vehicle}", space_after=8)
    para(f"Date: {sold_on}", space_after=0)
    para(f"Time: {time_now}", space_after=10)
    para("Dear Sirs,", space_after=8)
    body = para(
        f"I, {purchaser_name} sign to confirm that I have purchased the vehicle {reg} on the "
        f"{sold_on} from the seller {SELLER_PRIMARY} and therefore release all liability for this "
        f"vehicle from {SELLER_FULL_DOC}, as a result I {purchaser_name} now accept full liability "
        f"for the vehicle mentioned above beginning the {sold_on} and confirm by signing below that "
        f"I am the new owner of the vehicle {reg}.",
        space_after=8,
    )
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para(
        f"I also confirm that I will remove the vehicle purchased by me {purchaser_name} within 5 days of purchase and "
        f"payment having been received by {SELLER_PRIMARY}/Nationwide Assist Ltd.",
        space_after=14,
    )
    sign_pair()

    # ---- Page 2: Receipt of Sale ----
    doc.add_page_break()
    para("RECEIPT OF SALE:", bold=True, size=14, space_after=12)
    para(f"Payee: {purchaser_name}", space_after=0)
    para(f"Date: {sold_on}", space_after=12)

    products = doc.add_table(rows=2, cols=2)
    products.style = "Table Grid"
    products.columns[0].width = Inches(1.2)
    products.rows[0].cells[0].paragraphs[0].add_run("INV NO.").bold = True
    products.rows[0].cells[1].paragraphs[0].add_run("Products").bold = True
    products.rows[1].cells[0].text = ""
    product_cell = products.rows[1].cells[1]
    product_cell.paragraphs[0].add_run(product_name)
    product_cell.add_paragraph(f"Sale of vehicle registration {reg}.")

    para("", space_after=2)
    totals = doc.add_table(rows=3, cols=2)
    totals.style = "Table Grid"
    totals.columns[0].width = Inches(3.8)
    totals.columns[1].width = Inches(1.5)
    for i, (label, amount) in enumerate(
        [("Sub Total:", money_exc), ("VAT @ 20%", money_vat), ("Total:", money_inc)]
    ):
        label_p = totals.rows[i].cells[0].paragraphs[0]
        label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_p.add_run(label).bold = True
        totals.rows[i].cells[1].paragraphs[0].add_run(amount)

    para("", space_after=14)
    sign_pair()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
